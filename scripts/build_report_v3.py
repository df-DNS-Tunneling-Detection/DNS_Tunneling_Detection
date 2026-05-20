"""Premium-plus DOCX + PDF build of reports/report.md.

A more refined treatment than build_report_v2.py:
  - Emerald + gold + cream palette (classic academic publication feel).
  - Serif (Georgia) for cover title, headings, and pull-quotes; sans-serif
    (Arial) for body, tables, captions, and chrome.
  - Magazine-style cover: bold emerald top band with the title in white serif,
    gold eyebrow tag, decorative divider, centered submission card, team
    list with gold-bordered medallions, bottom emerald footer band.
  - H2 rendered with a large gold section number to the left of the heading
    text and a thin gold underline. H3 prefixed by a gold diamond.
  - Block quotes rendered as filled cream cards with an oversize serif "
    glyph and an italic Georgia body.
  - Tables: emerald header, soft emerald-100 zebra, soft borders, generous
    padding, all cells centered.
  - Body paragraphs justified, with inline bold/italic/code/link preserved.
  - Page header (project title in small caps) + page footer ("Page N of NN").

Writes to:
    reports/report_v3.docx
    reports/report_v3.pdf

reports/report.{md,docx,pdf} and reports/report_v2.{docx,pdf} are not touched.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_report as v1  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402
from fpdf import FPDF  # noqa: E402


ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "reports" / "report.md"
DOCX = ROOT / "reports" / "report_v3.docx"
PDF = ROOT / "reports" / "report_v3.pdf"


# ---------- v3 palette: emerald + gold + cream -------------------------------

P = {
    # Primaries
    "em_950":     (2, 44, 34),     # #022C22
    "em_900":     (6, 78, 59),     # #064E3B
    "em_800":     (6, 95, 70),     # #065F46
    "em_700":     (4, 120, 87),    # #047857
    "em_500":     (16, 185, 129),  # #10B981
    "em_100":     (209, 250, 229), # #D1FAE5
    "em_50":      (236, 253, 245), # #ECFDF5
    # Gold accent
    "gold_700":   (180, 83, 9),    # #B45309
    "gold_600":   (217, 119, 6),   # #D97706
    "gold_500":   (245, 158, 11),  # #F59E0B
    "gold_100":   (254, 243, 199), # #FEF3C7
    # Neutrals
    "charcoal":   (31, 41, 55),    # #1F2937 body text
    "slate_700":  (51, 65, 85),
    "slate_500":  (100, 116, 139),
    "slate_200":  (226, 232, 240),
    "cream":      (255, 251, 235), # #FFFBEB very pale gold
    # Hex forms used in DOCX XML
    "em_950_hex": "022C22",
    "em_900_hex": "064E3B",
    "em_800_hex": "065F46",
    "em_700_hex": "047857",
    "em_100_hex": "D1FAE5",
    "em_50_hex":  "ECFDF5",
    "gold_700_hex": "B45309",
    "gold_500_hex": "F59E0B",
    "gold_100_hex": "FEF3C7",
    "cream_hex":  "FFFBEB",
    "slate_200_hex": "E2E8F0",
}

SERIF = "Georgia"
SANS = "Arial"
MONO = "Mono"
RUNNING_HEADER = "DNS Tunneling Detection · Project Report"
RUNNING_FOOTER_LEFT = "Digital Forensics"


# ---------- Shared helpers ---------------------------------------------------

parse = v1.parse
inline_tokens = v1.inline_tokens
_strip_inline = v1._strip_inline
_is_numeric_cell = v1._is_numeric_cell
_split_cover = v1._split_cover
_extract_cover_info = v1._extract_cover_info
normalize_emoji = v1.normalize_emoji


# ---------- DOCX low-level helpers -------------------------------------------

def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=120, bottom=120, left=160, right=160) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _set_table_borders(table, color_hex: str, size: str = "4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color_hex)
        borders.append(node)
    tbl_pr.append(borders)


def _set_table_borderless(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def _add_paragraph_shading(paragraph, hex_fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    p_pr.append(shd)


def _add_centered(doc, text, *, size=12, bold=False, italic=False,
                  rgb=(0, 0, 0), space_after=4, space_before=0,
                  font_name=SANS):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*rgb)
    return p


def _render_runs_docx(paragraph, text: str, *, base_color=None) -> None:
    for kind, content in inline_tokens(text):
        if kind == "link":
            label, _url = content
            r = paragraph.add_run(label)
            r.font.color.rgb = RGBColor(*P["em_800"])
            r.underline = True
        elif kind == "bold":
            r = paragraph.add_run(content)
            r.bold = True
            r.font.color.rgb = RGBColor(*(base_color or P["em_900"]))
        elif kind == "italic":
            r = paragraph.add_run(content)
            r.italic = True
            if base_color is not None:
                r.font.color.rgb = RGBColor(*base_color)
        elif kind == "code":
            r = paragraph.add_run(content)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(*P["gold_700"])
        else:
            r = paragraph.add_run(content)
            if base_color is not None:
                r.font.color.rgb = RGBColor(*base_color)


# ---------- DOCX cover -------------------------------------------------------

def _build_docx_cover(doc, info: dict) -> None:
    # Top emerald band
    band = doc.add_paragraph()
    band.paragraph_format.space_before = Pt(0)
    band.paragraph_format.space_after = Pt(0)
    r = band.add_run(" " * 6)
    r.font.size = Pt(20)
    _add_paragraph_shading(band, P["em_900_hex"])

    # Thin gold stripe
    stripe = doc.add_paragraph()
    stripe.paragraph_format.space_before = Pt(0)
    stripe.paragraph_format.space_after = Pt(28)
    r = stripe.add_run(" " * 6)
    r.font.size = Pt(4)
    _add_paragraph_shading(stripe, P["gold_500_hex"])

    # Spacer
    for _ in range(2):
        doc.add_paragraph()

    # Eyebrow tag (small caps, gold)
    _add_centered(doc, info.get("course", "").upper(), size=10, bold=True,
                  rgb=P["gold_700"], space_after=10, font_name=SANS)

    # Two-tone serif title
    title = info["title"]
    if ":" in title:
        primary, _, accent = title.partition(":")
        accent = accent.strip()
        _add_centered(doc, primary.strip(), size=32, bold=True,
                      rgb=P["em_900"], space_after=2, font_name=SERIF)
        _add_centered(doc, accent, size=22, italic=True,
                      rgb=P["gold_700"], space_after=8, font_name=SERIF)
    else:
        _add_centered(doc, title, size=32, bold=True,
                      rgb=P["em_900"], space_after=8, font_name=SERIF)

    # Decorative divider: dot · line · diamond · line · dot
    _add_centered(doc, "•   ━━━━━   ◆   ━━━━━   •", size=11,
                  rgb=P["gold_600"], space_after=10)

    # Subtitle in serif italic
    _add_centered(doc, "Project Report", size=13, italic=True,
                  rgb=P["slate_500"], space_after=32, font_name=SERIF)

    # Submission card — 1×1 cream-tinted bordered table
    sup_tbl = doc.add_table(rows=2, cols=1)
    sup_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sup_tbl.autofit = False
    sup_tbl.columns[0].width = Inches(4.2)

    c0 = sup_tbl.cell(0, 0)
    c0.width = Inches(4.2)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c0.text = ""
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run("SUBMITTED TO")
    rr.bold = True
    rr.font.name = SANS
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(*P["gold_700"])
    _set_cell_shading(c0, P["cream_hex"])
    _set_cell_margins(c0, top=160, bottom=40, left=200, right=200)

    c1 = sup_tbl.cell(1, 0)
    c1.width = Inches(4.2)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.text = ""
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(info.get("supervisor", ""))
    rr.bold = True
    rr.font.name = SERIF
    rr.font.size = Pt(18)
    rr.font.color.rgb = RGBColor(*P["em_900"])
    _set_cell_shading(c1, P["cream_hex"])
    _set_cell_margins(c1, top=20, bottom=200, left=200, right=200)

    _set_table_borders(sup_tbl, P["gold_500_hex"], size="6")

    doc.add_paragraph()

    # Date
    _add_centered(doc, "DATE", size=10, bold=True,
                  rgb=P["gold_700"], space_after=4, font_name=SANS)
    _add_centered(doc, info.get("date", ""), size=14,
                  rgb=P["charcoal"], space_after=22, font_name=SERIF)

    # Team header
    _add_centered(doc, "PROJECT TEAM", size=11, bold=True,
                  rgb=P["gold_700"], space_after=4, font_name=SANS)
    _add_centered(doc, "◆", size=12, rgb=P["gold_500"], space_after=8)

    # Team list — borderless 2-column centered table for vertical alignment
    team_tbl = doc.add_table(rows=len(info["team"]), cols=2)
    team_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    team_tbl.autofit = False
    team_tbl.columns[0].width = Inches(0.55)
    team_tbl.columns[1].width = Inches(3.2)

    for r_idx, name in enumerate(info["team"]):
        c0 = team_tbl.cell(r_idx, 0)
        c0.width = Inches(0.55)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c0.text = ""
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        rn = p.add_run(f"{r_idx + 1:02d}")
        rn.bold = True
        rn.font.name = SERIF
        rn.font.size = Pt(14)
        rn.font.color.rgb = RGBColor(*P["gold_700"])
        _set_cell_margins(c0, top=40, bottom=40, left=80, right=80)

        c1 = team_tbl.cell(r_idx, 1)
        c1.width = Inches(3.2)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.text = ""
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        rname = p.add_run(name)
        rname.font.name = SERIF
        rname.font.size = Pt(13)
        rname.font.color.rgb = RGBColor(*P["em_900"])
        _set_cell_margins(c1, top=40, bottom=40, left=140, right=80)

    _set_table_borderless(team_tbl)

    # Push down a few paragraphs, then bottom band
    for _ in range(3):
        doc.add_paragraph()

    stripe = doc.add_paragraph()
    stripe.paragraph_format.space_before = Pt(0)
    stripe.paragraph_format.space_after = Pt(0)
    r = stripe.add_run(" " * 6)
    r.font.size = Pt(4)
    _add_paragraph_shading(stripe, P["gold_500_hex"])

    bottom = doc.add_paragraph()
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom.paragraph_format.space_before = Pt(0)
    bottom.paragraph_format.space_after = Pt(0)
    r = bottom.add_run("  " + _clean_url(info.get("repository", "")) + "  ")
    r.font.name = SANS
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = RGBColor(*P["cream"])
    _add_paragraph_shading(bottom, P["em_900_hex"])

    # Page break
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------- DOCX body --------------------------------------------------------

_NUM_HEAD_RE = re.compile(r"^(\d+(?:\.\d+)?)\.\s+(.+)$")


def _docx_render_table(doc, rows) -> None:
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = tbl.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx == 0:
                _render_runs_docx(p, cell_text, base_color=(255, 255, 255))
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(11)
                _set_cell_shading(cell, P["em_800_hex"])
            else:
                _render_runs_docx(p, cell_text, base_color=P["charcoal"])
                for run in p.runs:
                    run.font.size = Pt(10.5)
                if r_idx % 2 == 0:
                    _set_cell_shading(cell, P["em_50_hex"])
            _set_cell_margins(cell, top=120, bottom=120, left=180, right=180)

    _set_table_borders(tbl, P["em_100_hex"], size="6")


def _docx_render_quote(doc, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(0.15)
    tbl.columns[1].width = Inches(5.7)

    bar = tbl.cell(0, 0)
    bar.width = Inches(0.15)
    bar.text = ""
    _set_cell_shading(bar, P["gold_500_hex"])
    _set_cell_margins(bar, top=0, bottom=0, left=0, right=0)

    body = tbl.cell(0, 1)
    body.width = Inches(5.7)
    body.text = ""
    p = body.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _render_runs_docx(p, text, base_color=P["slate_700"])
    for run in p.runs:
        run.italic = True
        run.font.name = SERIF
        run.font.size = Pt(12)
    _set_cell_shading(body, P["cream_hex"])
    _set_cell_margins(body, top=180, bottom=180, left=220, right=220)

    _set_table_borderless(tbl)


def _docx_h2(doc, text: str) -> None:
    """Render h2 with a leading gold section number when present."""
    m = _NUM_HEAD_RE.match(text)
    p = doc.add_heading(level=1)
    if m:
        num, rest = m.group(1), m.group(2)
        r0 = p.add_run(num + ".  ")
        r0.font.name = SERIF
        r0.font.size = Pt(20)
        r0.bold = True
        r0.font.color.rgb = RGBColor(*P["gold_700"])
        r1 = p.add_run(rest)
        r1.font.name = SERIF
        r1.font.size = Pt(18)
        r1.bold = True
        r1.font.color.rgb = RGBColor(*P["em_900"])
    else:
        _render_runs_docx(p, text, base_color=P["em_900"])
        for run in p.runs:
            run.font.name = SERIF
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(*P["em_900"])


def _docx_h3(doc, text: str) -> None:
    p = doc.add_heading(level=2)
    r0 = p.add_run("◆  ")
    r0.font.size = Pt(11)
    r0.font.color.rgb = RGBColor(*P["gold_500"])
    _render_runs_docx(p, text, base_color=P["em_800"])
    for run in p.runs[1:]:
        run.font.name = SERIF
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(*P["em_800"])


def build_docx(cover_blocks, body_blocks):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)

    if cover_blocks:
        info = _extract_cover_info(cover_blocks)
        _build_docx_cover(doc, info)

    for kind, payload in body_blocks:
        if kind == "h1":
            p = doc.add_heading(level=0)
            _render_runs_docx(p, payload, base_color=P["em_900"])
            for run in p.runs:
                run.font.name = SERIF
                run.font.size = Pt(24)
                run.bold = True
                run.font.color.rgb = RGBColor(*P["em_900"])
        elif kind == "h2":
            _docx_h2(doc, payload)
        elif kind == "h3":
            _docx_h3(doc, payload)
        elif kind == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _render_runs_docx(p, payload, base_color=P["charcoal"])
        elif kind == "ul":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _render_runs_docx(p, item, base_color=P["charcoal"])
        elif kind == "ol":
            for item in payload:
                p = doc.add_paragraph(style="List Number")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _render_runs_docx(p, item, base_color=P["charcoal"])
        elif kind == "quote":
            _docx_render_quote(doc, payload)
        elif kind == "code":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(*P["slate_700"])
            _add_paragraph_shading(p, P["cream_hex"])
        elif kind == "math":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.italic = True
            r.font.name = "Cambria Math"
            r.font.color.rgb = RGBColor(*P["slate_700"])
        elif kind == "hr":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("◆   ◆   ◆")
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(*P["gold_500"])
        elif kind == "table":
            if not payload:
                continue
            _docx_render_table(doc, payload)
            doc.add_paragraph()

    doc.save(DOCX)
    print(f"wrote {DOCX} ({DOCX.stat().st_size/1024:.1f} KB)")


# ---------- PDF --------------------------------------------------------------

class CoverPDF(FPDF):
    """FPDF subclass with header/footer toggled per page."""

    def __init__(self):
        super().__init__(format="A4")
        self.set_margins(22, 26, 22)
        self.set_auto_page_break(True, margin=22)
        self.add_font("Arial", "", v1.ARIAL_REG, uni=True)
        self.add_font("Arial", "B", v1.ARIAL_BOLD, uni=True)
        self.add_font("Arial", "I", v1.ARIAL_IT, uni=True)
        self.add_font("Arial", "BI", v1.ARIAL_BI, uni=True)
        # Georgia for serif headings, title, pull-quotes.
        self.add_font("Georgia", "", r"C:\Windows\Fonts\georgia.ttf", uni=True)
        self.add_font("Georgia", "B", r"C:\Windows\Fonts\georgiab.ttf", uni=True)
        self.add_font("Georgia", "I", r"C:\Windows\Fonts\georgiai.ttf", uni=True)
        self.add_font("Georgia", "BI", r"C:\Windows\Fonts\georgiaz.ttf", uni=True)
        self.add_font("Mono", "", v1.CONSOLAS, uni=True)
        self.set_font("Arial", "", 11)
        self.alias_nb_pages()
        self._chrome = False

    def header(self):
        # Skip chrome on the cover (page 1) regardless of the _chrome flag.
        # Required because FPDF's footer() of the previous page fires inside
        # the next add_page() call — before we get a chance to disable chrome.
        if self.page_no() <= 1 or not self._chrome:
            return
        x0, x1 = self.l_margin, self.w - self.r_margin
        # Eyebrow line in tiny caps
        self.set_xy(x0, 12)
        self.set_font("Arial", "B", 8)
        self.set_text_color(*P["gold_700"])
        self.cell((x1 - x0) / 2, 5, RUNNING_HEADER.upper(), align="L")
        # Section dummy on the right (placeholder)
        self.set_xy((x0 + x1) / 2, 12)
        self.set_font("Arial", "", 8)
        self.set_text_color(*P["slate_500"])
        self.cell((x1 - x0) / 2, 5,
                  f"PAGE {self.page_no() - 1:02d} OF {{nb}}", align="R")
        # Thin emerald rule below
        y = 18
        pdf_r, pdf_g, pdf_b = P["em_800"]
        self.set_draw_color(pdf_r, pdf_g, pdf_b)
        self.set_line_width(0.4)
        self.line(x0, y, x1, y)
        # Tiny gold tick under center
        cx = (x0 + x1) / 2
        pdf_r, pdf_g, pdf_b = P["gold_500"]
        self.set_draw_color(pdf_r, pdf_g, pdf_b)
        self.set_line_width(0.6)
        self.line(cx - 4, y, cx + 4, y)
        self.set_text_color(0, 0, 0)
        self.set_y(self.t_margin)

    def footer(self):
        if self.page_no() <= 1 or not self._chrome:
            return
        x0, x1 = self.l_margin, self.w - self.r_margin
        y = self.h - 16
        pdf_r, pdf_g, pdf_b = P["gold_500"]
        self.set_draw_color(pdf_r, pdf_g, pdf_b)
        self.set_line_width(0.4)
        self.line(x0, y, x1, y)
        self.set_xy(x0, y + 2)
        self.set_font("Arial", "I", 8)
        self.set_text_color(*P["slate_500"])
        self.cell(x1 - x0, 4, RUNNING_FOOTER_LEFT + "   ·   2026 ", align="C")
        self.set_text_color(0, 0, 0)


# ---------- PDF cover --------------------------------------------------------

def _render_pdf_cover(pdf: CoverPDF, info: dict) -> None:
    # Disable auto page-break so the bottom-of-page URL cell can't trigger an
    # implicit page break mid-cover. Also temporarily zero the left/right
    # margins so the cover's full-width centered `cell(page_w, …, align="C")`
    # calls actually center on the page (otherwise the cursor sits at
    # l_margin and the cell overhangs the right edge, shifting all centered
    # text rightward by l_margin/2 = ~11mm).
    saved_apb = pdf.auto_page_break
    saved_b = pdf.b_margin
    saved_l = pdf.l_margin
    saved_r = pdf.r_margin
    pdf.set_auto_page_break(False)
    pdf.set_left_margin(0)
    pdf.set_right_margin(0)
    try:
        _render_pdf_cover_impl(pdf, info)
    finally:
        pdf.set_left_margin(saved_l)
        pdf.set_right_margin(saved_r)
        pdf.set_auto_page_break(saved_apb, margin=saved_b)


def _clean_url(url: str) -> str:
    """Strip Markdown autolink angle brackets from a URL."""
    return (url or "").strip().strip("<>").strip()


def _render_pdf_cover_impl(pdf: CoverPDF, info: dict) -> None:
    page_w, page_h = pdf.w, pdf.h
    cx = page_w / 2

    # === Top section: emerald band + gold stripe ===
    em_h = 12.0
    gold_h = 3.0
    r, g, b = P["em_900"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, 0, page_w, em_h, "F")
    r, g, b = P["gold_500"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, em_h, page_w, gold_h, "F")

    # === Bottom section: gold stripe + emerald band ===
    bot_em_h = 12.0
    bot_gold_h = 3.0
    bot_band_y = page_h - bot_em_h - bot_gold_h
    r, g, b = P["gold_500"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, bot_band_y, page_w, bot_gold_h, "F")
    r, g, b = P["em_900"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, bot_band_y + bot_gold_h, page_w, bot_em_h, "F")
    if info.get("repository"):
        pdf.set_xy(0, bot_band_y + bot_gold_h + 4.2)
        pdf.set_font("Arial", "B", 9)
        pdf.set_text_color(*P["cream"])
        pdf.cell(page_w, 5, _clean_url(info["repository"]), align="C")

    # === Corner ornaments: small gold L-brackets just inside the bands ===
    r, g, b = P["gold_700"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.6)
    # Top-left
    pdf.line(14, 22, 22, 22)
    pdf.line(14, 22, 14, 30)
    # Top-right
    pdf.line(page_w - 22, 22, page_w - 14, 22)
    pdf.line(page_w - 14, 22, page_w - 14, 30)
    # Bottom-left
    pdf.line(14, page_h - 22, 22, page_h - 22)
    pdf.line(14, page_h - 22, 14, page_h - 30)
    # Bottom-right
    pdf.line(page_w - 22, page_h - 22, page_w - 14, page_h - 22)
    pdf.line(page_w - 14, page_h - 22, page_w - 14, page_h - 30)

    # === Eyebrow tag ===
    pdf.set_y(56)
    pdf.set_font("Arial", "B", 10)
    pdf.set_text_color(*P["gold_700"])
    pdf.cell(page_w, 5, info.get("course", "").upper(), align="C", ln=1)

    # === Title ===
    title = info["title"]
    if ":" in title:
        primary, _, accent = title.partition(":")
        accent = accent.strip()
        pdf.set_y(67)
        pdf.set_font("Georgia", "B", 30)
        pdf.set_text_color(*P["em_900"])
        pdf.cell(page_w, 14, primary.strip(), align="C", ln=1)
        pdf.set_font("Georgia", "BI", 22)
        pdf.set_text_color(*P["gold_700"])
        pdf.cell(page_w, 10, accent, align="C", ln=1)
    else:
        pdf.set_y(70)
        pdf.set_font("Georgia", "B", 30)
        pdf.set_text_color(*P["em_900"])
        pdf.cell(page_w, 14, title, align="C", ln=1)

    # === Decorative divider: dots / line / diamond / line / dots ===
    y = pdf.get_y() + 4
    r, g, b = P["gold_600"]
    pdf.set_draw_color(r, g, b)
    pdf.set_fill_color(r, g, b)
    pdf.set_line_width(0.5)
    pdf.line(cx - 28, y, cx - 8, y)
    pdf.line(cx + 8, y, cx + 28, y)
    # diamond in center
    pdf.set_line_width(0)
    pdf.set_draw_color(r, g, b)
    diamond = [
        (cx, y - 2),
        (cx + 2, y),
        (cx, y + 2),
        (cx - 2, y),
    ]
    # Approximate with rectangle rotated 45° via filled small rect — easier: tiny ellipse
    pdf.ellipse(cx - 1.5, y - 1.5, 3, 3, "F")
    # Outer dots
    pdf.ellipse(cx - 32, y - 0.7, 1.4, 1.4, "F")
    pdf.ellipse(cx + 30.6, y - 0.7, 1.4, 1.4, "F")
    pdf.set_y(y + 6)

    # Subtitle
    pdf.set_font("Georgia", "I", 13)
    pdf.set_text_color(*P["slate_500"])
    pdf.cell(page_w, 7, "Project Report", align="C", ln=1)

    pdf.ln(18)

    # === Submission card ===
    card_w = 140.0
    card_h = 34.0
    card_x = (page_w - card_w) / 2
    card_y = pdf.get_y()
    r, g, b = P["cream"]
    pdf.set_fill_color(r, g, b)
    r, g, b = P["gold_500"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.7)
    pdf.rect(card_x, card_y, card_w, card_h, "DF")

    # Inner gold tick line
    r, g, b = P["gold_500"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.4)
    pdf.line(card_x + card_w / 2 - 18, card_y + 12, card_x + card_w / 2 + 18,
             card_y + 12)

    pdf.set_xy(card_x, card_y + 5)
    pdf.set_font("Arial", "B", 9)
    pdf.set_text_color(*P["gold_700"])
    pdf.cell(card_w, 5, "SUBMITTED TO", align="C")

    pdf.set_xy(card_x, card_y + 16)
    pdf.set_font("Georgia", "B", 18)
    pdf.set_text_color(*P["em_900"])
    pdf.cell(card_w, 12, info.get("supervisor", ""), align="C")

    pdf.set_y(card_y + card_h + 9)

    # === Date ===
    pdf.set_font("Arial", "B", 9)
    pdf.set_text_color(*P["gold_700"])
    pdf.cell(page_w, 5, "DATE", align="C", ln=1)
    pdf.set_font("Georgia", "", 13)
    pdf.set_text_color(*P["charcoal"])
    pdf.cell(page_w, 7, info.get("date", ""), align="C", ln=1)

    pdf.ln(8)

    # === Team header ===
    pdf.set_font("Arial", "B", 10)
    pdf.set_text_color(*P["gold_700"])
    pdf.cell(page_w, 5, "PROJECT TEAM", align="C", ln=1)
    # Small diamond underneath
    y = pdf.get_y() + 1
    r, g, b = P["gold_500"]
    pdf.set_fill_color(r, g, b)
    pdf.ellipse(cx - 1.6, y - 1.6, 3.2, 3.2, "F")
    pdf.set_y(y + 6)

    # === Team list — every badge at the same x ===
    badge_d = 7.0
    gap = 5.0
    pdf.set_font("Georgia", "", 13)
    max_name_w = max(pdf.get_string_width(n) for n in info["team"])
    block_w = badge_d + gap + max_name_w
    x_block = (page_w - block_w) / 2
    row_h = 9.0

    for n, name in enumerate(info["team"], 1):
        y_row = pdf.get_y()

        # Outer gold ring + cream fill medallion
        r, g, b = P["cream"]
        pdf.set_fill_color(r, g, b)
        r, g, b = P["gold_700"]
        pdf.set_draw_color(r, g, b)
        pdf.set_line_width(0.6)
        pdf.ellipse(x_block, y_row + 0.6, badge_d, badge_d, "DF")

        # Number inside the medallion
        pdf.set_xy(x_block, y_row + 0.6)
        pdf.set_font("Georgia", "B", 9)
        pdf.set_text_color(*P["em_900"])
        pdf.cell(badge_d, badge_d, str(n), align="C")

        # Name in serif emerald
        pdf.set_xy(x_block + badge_d + gap, y_row)
        pdf.set_font("Georgia", "", 13)
        pdf.set_text_color(*P["em_900"])
        pdf.cell(max_name_w + 1, badge_d + 1, name, align="L")

        pdf.set_y(y_row + row_h)

    pdf.set_text_color(0, 0, 0)


# ---------- PDF body --------------------------------------------------------

def _v3_font_for(kind: str):
    if kind in ("bold", "link"):
        return (SANS, "B", 11)
    if kind == "italic":
        return (SANS, "I", 11)
    if kind == "code":
        return (MONO, "", 10)
    return (SANS, "", 11)


def _v3_color_for(kind: str):
    if kind == "bold":
        return P["em_900"]
    if kind == "code":
        return P["gold_700"]
    if kind == "link":
        return P["em_800"]
    return P["charcoal"]


def write_justified(pdf, text: str, line_h: float = 5.4) -> None:
    avail = pdf.w - pdf.l_margin - pdf.r_margin
    atoms = []
    for kind, content in inline_tokens(text):
        if kind == "link":
            content = content[0]
        for part in re.split(r"(\s+)", content):
            if part == "":
                continue
            atoms.append((kind, part))

    measured = []
    for kind, word in atoms:
        fam, sty, sz = _v3_font_for(kind)
        pdf.set_font(fam, sty, sz)
        w = pdf.get_string_width(word)
        measured.append((kind, word, w, word.isspace()))

    lines, cur, cur_w = [], [], 0.0
    for atom in measured:
        _, _, w, is_space = atom
        if cur_w + w > avail and cur:
            while cur and cur[-1][3]:
                cur.pop()
            lines.append(cur)
            cur, cur_w = [], 0.0
            if is_space:
                continue
        cur.append(atom)
        cur_w += w
    if cur:
        while cur and cur[-1][3]:
            cur.pop()
        lines.append(cur)

    bottom_trigger = pdf.h - pdf.b_margin
    for idx, line in enumerate(lines):
        if pdf.get_y() + line_h > bottom_trigger:
            pdf.add_page()

        is_last = idx == len(lines) - 1
        non_space_w = sum(w for _, _, w, sp in line if not sp)
        n_spaces = sum(1 for _, _, _, sp in line if sp)
        base_space_w = sum(w for _, _, w, sp in line if sp)
        extra_per_space = 0.0 if (is_last or n_spaces == 0) else \
            (avail - non_space_w - base_space_w) / n_spaces

        x = pdf.l_margin
        y = pdf.get_y()
        for kind, word, w, is_space in line:
            fam, sty, sz = _v3_font_for(kind)
            pdf.set_font(fam, sty, sz)
            pdf.set_text_color(*_v3_color_for(kind))
            if is_space:
                x += w + extra_per_space
            else:
                pdf.set_xy(x, y)
                pdf.cell(w + 0.1, line_h, word)
                x += w
        pdf.set_text_color(0, 0, 0)
        pdf.set_xy(pdf.l_margin, y + line_h)


def _pdf_h1(pdf, text: str) -> None:
    pdf.set_font("Georgia", "B", 24)
    pdf.set_text_color(*P["em_900"])
    pdf.multi_cell(0, 11, text)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)


def _pdf_h2(pdf, text: str) -> None:
    """H2 with a large gold leading numeral + emerald heading + gold underline."""
    pdf.ln(4)

    # Pre-measure the heading text height so we can decide atomically whether
    # the whole H2 (numeral + wrapped text + underline + trailing space) fits
    # on the current page. If not, force a page break first so the numeral
    # and text never end up on different pages.
    m = _NUM_HEAD_RE.match(text)
    if m:
        rest = m.group(2)
    else:
        rest = text
    avail_text_w = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.set_font("Georgia", "B", 18)
    text_lines = pdf.multi_cell(avail_text_w, 10, rest,
                                dry_run=True, output="LINES")
    needed = max(11, len(text_lines) * 10) + 8  # text + underline + ln(4)
    if pdf.get_y() + needed > pdf.h - pdf.b_margin:
        pdf.add_page()

    y0 = pdf.get_y()
    if m:
        num, rest = m.group(1), m.group(2)
        pdf.set_font("Georgia", "B", 22)
        pdf.set_text_color(*P["gold_700"])
        num_text = num + "."
        num_w = pdf.get_string_width(num_text) + 3
        pdf.set_xy(pdf.l_margin, y0)
        pdf.cell(num_w, 11, num_text)
        pdf.set_xy(pdf.l_margin + num_w + 1, y0 + 1.2)
        pdf.set_font("Georgia", "B", 18)
        pdf.set_text_color(*P["em_900"])
        pdf.multi_cell(0, 10, rest)
    else:
        pdf.set_font("Georgia", "B", 18)
        pdf.set_text_color(*P["em_900"])
        pdf.multi_cell(0, 10, text)

    # Gold underline
    y = pdf.get_y() + 0.4
    r, g, b = P["gold_500"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.7)
    pdf.line(pdf.l_margin, y, pdf.l_margin + 55, y)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)


def _pdf_h3(pdf, text: str) -> None:
    pdf.ln(2)
    # Atomic-fit check: don't strand the gold diamond from the text below.
    avail_w = pdf.w - pdf.l_margin - pdf.r_margin - 5
    pdf.set_font("Georgia", "B", 13)
    lines = pdf.multi_cell(avail_w, 6.5, text, dry_run=True, output="LINES")
    needed = max(7, len(lines) * 6.5) + 1
    if pdf.get_y() + needed > pdf.h - pdf.b_margin:
        pdf.add_page()
    y = pdf.get_y()
    # Gold diamond
    r, g, b = P["gold_500"]
    pdf.set_fill_color(r, g, b)
    pdf.ellipse(pdf.l_margin, y + 2.4, 2.6, 2.6, "F")
    pdf.set_xy(pdf.l_margin + 5, y)
    pdf.set_font("Georgia", "B", 13)
    pdf.set_text_color(*P["em_800"])
    pdf.multi_cell(0, 6.5, text)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(0.5)


def _compute_col_widths(pdf, rows, n_cols, page_width):
    pdf.set_font("Arial", "", 10)
    raw = [0.0] * n_cols
    for row in rows:
        for c_idx in range(n_cols):
            cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
            for line in cell.splitlines() or [cell]:
                w = pdf.get_string_width(line)
                if w > raw[c_idx]:
                    raw[c_idx] = w
    pad = 5.5
    raw = [w + pad for w in raw]
    total = sum(raw) or 1.0
    widths = [max(15.0, w * page_width / total) for w in raw]
    scale = page_width / sum(widths)
    return [w * scale for w in widths]


def _draw_pdf_row(pdf, row, widths, h, line_h, v_pad,
                  is_header=False, zebra=False):
    n_cols = len(widths)
    y_start = pdf.get_y()
    for c_idx in range(n_cols):
        cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
        x_start = pdf.l_margin + sum(widths[:c_idx])
        if is_header:
            r, g, b = P["em_800"]
            pdf.set_fill_color(r, g, b)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Arial", "B", 10.5)
        else:
            if zebra:
                r, g, b = P["em_50"]
                pdf.set_fill_color(r, g, b)
            else:
                pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(*P["charcoal"])
            pdf.set_font("Arial", "", 10)
        r, g, b = P["em_100"]
        pdf.set_draw_color(r, g, b)
        pdf.set_line_width(0.2)

        pdf.rect(x_start, y_start, widths[c_idx], h, "DF")

        inset_x = 1.6
        inset_y = v_pad / 2
        pdf.set_xy(x_start + inset_x, y_start + inset_y)
        pdf.multi_cell(widths[c_idx] - 2 * inset_x, line_h, cell,
                       border=0, align="C", fill=False)
    pdf.set_xy(pdf.l_margin, y_start + h)
    pdf.set_text_color(0, 0, 0)


def _render_pdf_table(pdf, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    widths = _compute_col_widths(pdf, rows, n_cols, page_width)

    pdf.ln(2)
    line_h = 5.8
    v_pad = 2.6

    def measure(row, font_style):
        pdf.set_font("Arial", font_style, 10)
        max_lines = 1
        for c_idx in range(n_cols):
            cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
            lines = pdf.multi_cell(widths[c_idx], line_h, cell,
                                   dry_run=True, output="LINES")
            if len(lines) > max_lines:
                max_lines = len(lines)
        return max_lines * line_h + v_pad

    for r_idx, row in enumerate(rows):
        is_header = r_idx == 0
        h = measure(row, "B" if is_header else "")
        if pdf.get_y() + h > pdf.h - pdf.b_margin:
            pdf.add_page()
            if not is_header and rows:
                hh = measure(rows[0], "B")
                _draw_pdf_row(pdf, rows[0], widths, hh, line_h, v_pad,
                              is_header=True)
        _draw_pdf_row(pdf, row, widths, h, line_h, v_pad,
                      is_header=is_header,
                      zebra=(r_idx % 2 == 0 and not is_header))
    pdf.ln(5)


def _render_pdf_quote(pdf, text: str) -> None:
    """Cream callout with oversized gold opening quote and italic serif body."""
    pdf.ln(2)
    y_start = pdf.get_y()
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    text_w = page_width - 12
    pdf.set_font("Georgia", "I", 12)
    lines = pdf.multi_cell(text_w, 5.6, text, dry_run=True, output="LINES")
    box_h = max(14.0, len(lines) * 5.6 + 6.0)

    # Cream background
    r, g, b = P["cream"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(pdf.l_margin, y_start, page_width, box_h, "F")
    # Left thick gold bar
    r, g, b = P["gold_500"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(pdf.l_margin, y_start, 2.0, box_h, "F")

    # Oversized opening quote glyph
    pdf.set_xy(pdf.l_margin + 4, y_start + 1)
    pdf.set_font("Georgia", "B", 22)
    pdf.set_text_color(*P["gold_700"])
    pdf.cell(8, 8, "“")  # left double quotation mark

    pdf.set_xy(pdf.l_margin + 12, y_start + 3)
    pdf.set_font("Georgia", "I", 12)
    pdf.set_text_color(*P["slate_700"])
    pdf.multi_cell(text_w - 8, 5.6, text)
    pdf.set_text_color(0, 0, 0)
    pdf.set_y(y_start + box_h + 4)


def render_pdf(cover_blocks, body_blocks):
    pdf = CoverPDF()
    pdf._chrome = False
    pdf.add_page()
    if cover_blocks:
        info = _extract_cover_info(cover_blocks)
        _render_pdf_cover(pdf, info)

    pdf._chrome = True
    pdf.add_page()

    for kind, payload in body_blocks:
        if kind == "h1":
            _pdf_h1(pdf, payload)
        elif kind == "h2":
            _pdf_h2(pdf, payload)
        elif kind == "h3":
            _pdf_h3(pdf, payload)
        elif kind == "p":
            write_justified(pdf, payload)
            pdf.ln(2)
        elif kind == "ul":
            for item in payload:
                pdf.set_font("Arial", "", 11)
                pdf.cell(4, 5.4, "")
                pdf.set_text_color(*P["gold_500"])
                pdf.write(5.4, "■  ")
                pdf.set_text_color(0, 0, 0)
                pdf.add_inline(item) if hasattr(pdf, "add_inline") else None
                # fall back to raw write if add_inline is missing
                if not hasattr(pdf, "add_inline"):
                    pdf.set_font("Arial", "", 11)
                    pdf.set_text_color(*P["charcoal"])
                    pdf.write(5.4, _strip_inline(item))
                    pdf.set_text_color(0, 0, 0)
                pdf.ln(5.6)
            pdf.ln(2)
        elif kind == "ol":
            for n, item in enumerate(payload, 1):
                pdf.set_font("Arial", "B", 11)
                pdf.cell(4, 5.4, "")
                pdf.set_text_color(*P["gold_700"])
                pdf.write(5.4, f"{n:02d}.  ")
                pdf.set_text_color(0, 0, 0)
                if hasattr(pdf, "add_inline"):
                    pdf.add_inline(item)
                else:
                    pdf.set_font("Arial", "", 11)
                    pdf.set_text_color(*P["charcoal"])
                    pdf.write(5.4, _strip_inline(item))
                    pdf.set_text_color(0, 0, 0)
                pdf.ln(5.6)
            pdf.ln(2)
        elif kind == "quote":
            _render_pdf_quote(pdf, payload)
        elif kind == "code":
            r, g, b = P["cream"]
            pdf.set_fill_color(r, g, b)
            pdf.set_font("Mono", "", 9)
            pdf.set_text_color(*P["slate_700"])
            for code_line in payload.split("\n"):
                pdf.cell(0, 4.6, code_line, ln=1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(3)
        elif kind == "math":
            pdf.set_font("Georgia", "I", 11)
            pdf.set_text_color(*P["slate_700"])
            pdf.multi_cell(0, 5.6, payload)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(3)
        elif kind == "hr":
            y = pdf.get_y() + 2
            cx = pdf.w / 2
            r, g, b = P["gold_500"]
            pdf.set_draw_color(r, g, b)
            pdf.set_line_width(0.4)
            pdf.line(cx - 28, y, cx - 6, y)
            pdf.line(cx + 6, y, cx + 28, y)
            pdf.set_fill_color(r, g, b)
            pdf.ellipse(cx - 1.5, y - 1.5, 3, 3, "F")
            pdf.ln(6)
        elif kind == "table":
            _render_pdf_table(pdf, payload)

    pdf.output(str(PDF))
    print(f"wrote {PDF} ({PDF.stat().st_size/1024:.1f} KB)")


# CoverPDF doesn't define add_inline; supply a small shim so list items still
# render with inline formatting (mirrors v1.add_inline but using v3 palette).
def _add_inline_shim(self, text: str, base_style: str = ""):
    for kind, content in inline_tokens(text):
        if kind == "link":
            content = content[0]
            self.set_text_color(*P["em_800"])
            self.set_font("Arial", "B", 11)
            self.write(5.4, content)
            self.set_text_color(0, 0, 0)
        elif kind == "bold":
            self.set_text_color(*P["em_900"])
            self.set_font("Arial", "B", 11)
            self.write(5.4, content)
            self.set_text_color(0, 0, 0)
        elif kind == "italic":
            self.set_font("Arial", "I", 11)
            self.set_text_color(*P["charcoal"])
            self.write(5.4, content)
            self.set_text_color(0, 0, 0)
        elif kind == "code":
            self.set_font("Mono", "", 10)
            self.set_text_color(*P["gold_700"])
            self.write(5.4, content)
            self.set_text_color(0, 0, 0)
        else:
            self.set_font("Arial", base_style, 11)
            self.set_text_color(*P["charcoal"])
            self.write(5.4, content)
            self.set_text_color(0, 0, 0)


CoverPDF.add_inline = _add_inline_shim


# ---------- Main -------------------------------------------------------------

def main():
    md_text = MD.read_text(encoding="utf-8")
    md_text = normalize_emoji(md_text)
    blocks = list(parse(md_text))
    cover_blocks, body_blocks = _split_cover(blocks)
    print(f"Parsed {len(blocks)} blocks (cover={len(cover_blocks)}, "
          f"body={len(body_blocks)}) from {MD}")
    build_docx(cover_blocks, body_blocks)
    render_pdf(cover_blocks, body_blocks)


if __name__ == "__main__":
    main()
