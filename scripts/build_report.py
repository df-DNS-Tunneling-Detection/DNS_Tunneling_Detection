"""Convert reports/report.md into reports/report.docx and reports/report.pdf.

Self-contained — relies only on python-docx and fpdf2, both pure-Python.

Supports:
- Headings (#, ##, ###)
- Paragraphs with inline **bold**, *italic*, `code`, [link](url)
- Tables (| col | col |)
- Code blocks (``` ... ```)
- Bulleted (-) and numbered (1.) lists
- Blockquotes (> ...)
- Horizontal rules (---)
- Math blocks ($$ ... $$) and inline math ($ ... $) rendered as italic text
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "reports" / "report.md"
DOCX = ROOT / "reports" / "report.docx"
PDF = ROOT / "reports" / "report.pdf"

ARIAL_REG = r"C:\Windows\Fonts\arial.ttf"
ARIAL_BOLD = r"C:\Windows\Fonts\arialbd.ttf"
ARIAL_IT = r"C:\Windows\Fonts\ariali.ttf"
ARIAL_BI = r"C:\Windows\Fonts\arialbi.ttf"
CONSOLAS = r"C:\Windows\Fonts\consola.ttf"


# ---------- Palette -----------------------------------------------------------
# Single source of truth for colors. Teal primary + amber accent + warm-gray text.

PALETTE = {
    # Headings
    "h1_rgb":            (19, 78, 74),     # #134E4A dark teal
    "h2_rgb":            (15, 118, 110),   # #0F766E teal
    "h3_rgb":            (180, 83, 9),     # #B45309 amber
    # Inline
    "link_rgb":          (14, 116, 144),   # #0E7490 cyan
    # Tables
    "tbl_header_fill":   "0F766E",        # teal
    "tbl_header_hex":    "0F766E",
    "tbl_header_rgb":    (15, 118, 110),
    "tbl_zebra_fill":    "ECFEFF",        # very pale cyan
    "tbl_zebra_rgb":     (236, 254, 255),
    "tbl_border_hex":    "5EEAD4",        # soft teal
    "tbl_border_rgb":    (94, 234, 212),
    # Misc
    "quote_rgb":         (87, 83, 78),
    "code_fill_rgb":     (252, 248, 234), # warm cream
    "hr_rgb":            (180, 83, 9),    # amber rule
    # Cover page
    "cover_band_rgb":    (19, 78, 74),    # dark teal top band
    "cover_band_b_rgb":  (180, 83, 9),    # amber bottom band
    "cover_title_rgb":   (19, 78, 74),
    "cover_subtitle_rgb":(15, 118, 110),
    "cover_label_rgb":   (180, 83, 9),
    "cover_value_rgb":   (40, 40, 40),
    "cover_team_rgb":    (15, 118, 110),
}


# ---------- Markdown parsing ---------------------------------------------------

BLOCK_RE = {
    "h1": re.compile(r"^# (.+)"),
    "h2": re.compile(r"^## (.+)"),
    "h3": re.compile(r"^### (.+)"),
    "hr": re.compile(r"^---+$"),
    "code_fence": re.compile(r"^```(\w*)$"),
    "ul": re.compile(r"^(\s*)- (.+)"),
    "ol": re.compile(r"^(\s*)\d+\. (.+)"),
    "quote": re.compile(r"^> (.+)"),
    "table_row": re.compile(r"^\|.*\|\s*$"),
    "table_sep": re.compile(r"^\|[\s\-:|]+\|\s*$"),
    "math_block": re.compile(r"^\$\$"),
}


def parse(md_text: str):
    """Yield (kind, payload) blocks."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if BLOCK_RE["h1"].match(line):
            yield ("h1", BLOCK_RE["h1"].match(line).group(1))
            i += 1
        elif BLOCK_RE["h2"].match(line):
            yield ("h2", BLOCK_RE["h2"].match(line).group(1))
            i += 1
        elif BLOCK_RE["h3"].match(line):
            yield ("h3", BLOCK_RE["h3"].match(line).group(1))
            i += 1
        elif BLOCK_RE["hr"].match(line):
            yield ("hr", None)
            i += 1
        elif BLOCK_RE["code_fence"].match(line):
            buf = []
            i += 1
            while i < len(lines) and not BLOCK_RE["code_fence"].match(lines[i]):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            yield ("code", "\n".join(buf))
        elif BLOCK_RE["math_block"].match(line):
            stripped = line.strip()
            # Single-line case: $$ expr $$
            if stripped.endswith("$$") and len(stripped) > 4:
                yield ("math", stripped[2:-2].strip())
                i += 1
            else:
                buf = []
                i += 1
                while i < len(lines) and not BLOCK_RE["math_block"].match(lines[i]):
                    buf.append(lines[i])
                    i += 1
                i += 1
                yield ("math", "\n".join(buf))
        elif BLOCK_RE["table_row"].match(line):
            rows = []
            while i < len(lines) and BLOCK_RE["table_row"].match(lines[i]):
                if BLOCK_RE["table_sep"].match(lines[i]):
                    i += 1
                    continue
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            yield ("table", rows)
        elif BLOCK_RE["ul"].match(line):
            items = []
            while i < len(lines) and BLOCK_RE["ul"].match(lines[i]):
                items.append(BLOCK_RE["ul"].match(lines[i]).group(2))
                i += 1
            yield ("ul", items)
        elif BLOCK_RE["ol"].match(line):
            items = []
            while i < len(lines) and BLOCK_RE["ol"].match(lines[i]):
                items.append(BLOCK_RE["ol"].match(lines[i]).group(2))
                i += 1
            yield ("ol", items)
        elif BLOCK_RE["quote"].match(line):
            items = []
            while i < len(lines) and BLOCK_RE["quote"].match(lines[i]):
                items.append(BLOCK_RE["quote"].match(lines[i]).group(1))
                i += 1
            yield ("quote", " ".join(items))
        elif line.strip() == "":
            i += 1
        else:
            # Paragraph — gather consecutive non-empty, non-block lines
            buf = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not any(
                BLOCK_RE[k].match(lines[i])
                for k in ("h1", "h2", "h3", "hr", "code_fence", "ul", "ol", "quote", "table_row", "math_block")
            ):
                buf.append(lines[i])
                i += 1
            yield ("p", " ".join(buf))


# ---------- Inline parsing -----------------------------------------------------

# tokenizer: yields (style_set, text) where style_set is a frozenset of {"bold","italic","code","link"}
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)|"           # bold
    r"(`[^`]+`)|"                 # code
    r"(\*[^*\n]+\*)|"             # italic
    r"(\[[^\]]+\]\([^)]+\))|"     # link
    r"(\$[^$\n]+\$)"              # inline math (treated as italic)
)


def inline_tokens(text: str):
    """Yield (kind, content) tokens from inline-formatted text."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            yield ("plain", text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**"):
            yield ("bold", token[2:-2])
        elif token.startswith("`"):
            yield ("code", token[1:-1])
        elif token.startswith("*"):
            yield ("italic", token[1:-1])
        elif token.startswith("["):
            inner = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            yield ("link", (inner.group(1), inner.group(2)))
        elif token.startswith("$"):
            yield ("italic", token[1:-1])  # render math as italic
        pos = m.end()
    if pos < len(text):
        yield ("plain", text[pos:])


# ---------- DOCX rendering -----------------------------------------------------

def render_runs_docx(paragraph, text: str):
    for kind, content in inline_tokens(text):
        if kind == "link":
            label, _url = content
            r = paragraph.add_run(label)
            r.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
            r.underline = True
        elif kind == "bold":
            paragraph.add_run(content).bold = True
        elif kind == "italic":
            paragraph.add_run(content).italic = True
        elif kind == "code":
            r = paragraph.add_run(content)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            paragraph.add_run(content)


_NUMERIC_RE = re.compile(r"^[\s\-+]?\$?\d[\d,]*(\.\d+)?\s*%?$")


def _is_numeric_cell(text: str) -> bool:
    stripped = _strip_inline(text).strip()
    if not stripped:
        return False
    return bool(_NUMERIC_RE.match(stripped.replace(" ", "")))


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    """Cell padding in twentieths of a point (1/20 pt). 80 ≈ 4pt, 120 ≈ 6pt."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _set_table_borders(table, color="A6B0C0", size="6") -> None:
    """Light, uniform borders around every cell. size is in eighths of a point."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def _render_docx_table(doc, rows) -> None:
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = tbl.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            render_runs_docx(p, cell_text)

            if r_idx == 0:
                _set_cell_shading(cell, PALETTE["tbl_header_hex"])
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10.5)
            else:
                if r_idx % 2 == 0:
                    _set_cell_shading(cell, PALETTE["tbl_zebra_fill"])
                for run in p.runs:
                    run.font.size = Pt(10.5)
            _set_cell_margins(cell)

    _set_table_borders(tbl, color=PALETTE["tbl_border_hex"], size="6")


# ---------- Cover-page extraction --------------------------------------------

def _split_cover(blocks):
    """Split blocks into (cover_blocks, body_blocks) at the first hr after h1.

    If the document does not begin with an h1, no cover is extracted.
    """
    if not blocks or blocks[0][0] != "h1":
        return [], blocks
    for i, (kind, _) in enumerate(blocks):
        if i > 0 and kind == "hr":
            return blocks[:i], blocks[i + 1 :]
    return [], blocks


def _extract_cover_info(cover_blocks) -> dict:
    """Pull title, supervisor, course, date, repository, and team list out of
    the cover blocks parsed from the markdown.
    """
    info = {
        "title": "",
        "subtitle": "Project Report",
        "supervisor": "",
        "course": "",
        "date": "",
        "repository": "",
        "team": [],
    }
    label_map = {
        "submitted to": "supervisor",
        "course": "course",
        "date": "date",
        "repository": "repository",
    }
    for kind, payload in cover_blocks:
        if kind == "h1":
            info["title"] = payload
        elif kind == "table":
            rows = payload
            if not rows:
                continue
            header_l = [_strip_inline(c).strip().lower() for c in rows[0]]
            # First metadata table has columns Field | Value.
            if header_l[:2] == ["field", "value"]:
                for row in rows[1:]:
                    if len(row) < 2:
                        continue
                    label = _strip_inline(row[0]).strip().lower()
                    value = _strip_inline(row[1]).strip()
                    if label in label_map:
                        info[label_map[label]] = value
            elif "name" in header_l:
                name_col = header_l.index("name")
                for row in rows[1:]:
                    if len(row) > name_col:
                        info["team"].append(_strip_inline(row[name_col]).strip())
        elif kind == "ol":
            # Numbered list fallback for team.
            if not info["team"]:
                info["team"] = [_strip_inline(s).strip() for s in payload]
    return info


# ---------- DOCX cover page --------------------------------------------------

def _add_centered(doc, text, *, size=12, bold=False, italic=False,
                  rgb=(0, 0, 0), space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*rgb)
    return p


def _add_paragraph_shading(paragraph, hex_fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    p_pr.append(shd)


def _build_docx_cover(doc, info: dict) -> None:
    # Top color band
    band = doc.add_paragraph()
    band.paragraph_format.space_before = Pt(0)
    band.paragraph_format.space_after = Pt(24)
    band.add_run(" " * 4).font.size = Pt(14)
    _add_paragraph_shading(band, "134E4A")

    # Vertical space, then big centered title
    for _ in range(3):
        doc.add_paragraph()
    _add_centered(doc, info["title"], size=28, bold=True,
                  rgb=PALETTE["cover_title_rgb"], space_after=6)

    # Amber accent rule (a row of em-dashes, centered)
    _add_centered(doc, "─" * 20, size=14,
                  rgb=PALETTE["cover_band_b_rgb"], space_after=18)

    # Subtitle / course
    _add_centered(doc, info.get("course", ""), size=16, italic=True,
                  rgb=PALETTE["cover_subtitle_rgb"], space_after=4)
    _add_centered(doc, info["subtitle"], size=13, italic=True,
                  rgb=PALETTE["cover_subtitle_rgb"], space_after=28)

    # Supervisor block
    _add_centered(doc, "Submitted to", size=11, bold=True,
                  rgb=PALETTE["cover_label_rgb"], space_after=2)
    _add_centered(doc, info.get("supervisor", ""), size=14, bold=True,
                  rgb=PALETTE["cover_value_rgb"], space_after=18)

    # Date
    _add_centered(doc, "Date", size=11, bold=True,
                  rgb=PALETTE["cover_label_rgb"], space_after=2)
    _add_centered(doc, info.get("date", ""), size=13,
                  rgb=PALETTE["cover_value_rgb"], space_after=28)

    # Team
    _add_centered(doc, "Project Team", size=14, bold=True,
                  rgb=PALETTE["cover_team_rgb"], space_after=8)
    for n, name in enumerate(info["team"], 1):
        _add_centered(doc, f"{n}.   {name}", size=12,
                      rgb=PALETTE["cover_value_rgb"], space_after=3)

    # Bottom amber band
    for _ in range(2):
        doc.add_paragraph()
    band = doc.add_paragraph()
    band.add_run(" " * 4).font.size = Pt(10)
    _add_paragraph_shading(band, "B45309")

    # Page break to start body on a fresh page
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------- PDF cover page ---------------------------------------------------

def _render_pdf_cover(pdf, info: dict) -> None:
    saved_apb = pdf.auto_page_break
    saved_margin = pdf.b_margin
    pdf.set_auto_page_break(False)
    try:
        _render_pdf_cover_impl(pdf, info)
    finally:
        pdf.set_auto_page_break(saved_apb, margin=saved_margin)


def _render_pdf_cover_impl(pdf, info: dict) -> None:
    page_w = pdf.w
    page_h = pdf.h
    margin = pdf.l_margin

    # Top dark teal band
    band_h = 22.0
    r, g, b = PALETTE["cover_band_rgb"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, 0, page_w, band_h, "F")

    # Amber accent stripe just below
    r, g, b = PALETTE["cover_band_b_rgb"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, band_h, page_w, 3.0, "F")

    # Title block, centered around 30% of page
    y_title = 70.0
    pdf.set_xy(margin, y_title)
    pdf.set_font("Arial", "B", 26)
    pdf.set_text_color(*PALETTE["cover_title_rgb"])
    pdf.multi_cell(page_w - 2 * margin, 12, info["title"], align="C")

    # Amber rule under title
    y = pdf.get_y() + 4
    r, g, b = PALETTE["cover_band_b_rgb"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.8)
    cx = page_w / 2
    pdf.line(cx - 30, y, cx + 30, y)
    pdf.set_y(y + 8)

    # Course + subtitle
    pdf.set_font("Arial", "I", 15)
    pdf.set_text_color(*PALETTE["cover_subtitle_rgb"])
    pdf.cell(page_w - 2 * margin, 8, info.get("course", ""), align="C", ln=1)
    pdf.set_font("Arial", "I", 12)
    pdf.cell(page_w - 2 * margin, 7, info["subtitle"], align="C", ln=1)

    # Spacer
    pdf.ln(20)

    # Supervisor
    pdf.set_font("Arial", "B", 11)
    pdf.set_text_color(*PALETTE["cover_label_rgb"])
    pdf.cell(page_w - 2 * margin, 6, "Submitted to", align="C", ln=1)
    pdf.set_font("Arial", "B", 14)
    pdf.set_text_color(*PALETTE["cover_value_rgb"])
    pdf.cell(page_w - 2 * margin, 8, info.get("supervisor", ""), align="C", ln=1)
    pdf.ln(12)

    # Date
    pdf.set_font("Arial", "B", 11)
    pdf.set_text_color(*PALETTE["cover_label_rgb"])
    pdf.cell(page_w - 2 * margin, 6, "Date", align="C", ln=1)
    pdf.set_font("Arial", "", 12)
    pdf.set_text_color(*PALETTE["cover_value_rgb"])
    pdf.cell(page_w - 2 * margin, 7, info.get("date", ""), align="C", ln=1)
    pdf.ln(16)

    # Team
    pdf.set_font("Arial", "B", 14)
    pdf.set_text_color(*PALETTE["cover_team_rgb"])
    pdf.cell(page_w - 2 * margin, 8, "Project Team", align="C", ln=1)

    # Decorative short rule under "Project Team"
    y = pdf.get_y() + 2
    r, g, b = PALETTE["cover_team_rgb"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.4)
    pdf.line(cx - 18, y, cx + 18, y)
    pdf.set_y(y + 6)

    pdf.set_font("Arial", "", 12)
    pdf.set_text_color(*PALETTE["cover_value_rgb"])
    for n, name in enumerate(info["team"], 1):
        pdf.cell(page_w - 2 * margin, 7, f"{n}.   {name}", align="C", ln=1)

    # Bottom amber band
    bottom_band_h = 14.0
    r, g, b = PALETTE["cover_band_b_rgb"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, page_h - bottom_band_h, page_w, bottom_band_h, "F")

    # Repository URL inside the bottom band (small, white)
    if info.get("repository"):
        pdf.set_xy(margin, page_h - bottom_band_h + 4.2)
        pdf.set_font("Arial", "B", 9)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(page_w - 2 * margin, 5, info["repository"], align="C", ln=1)
    pdf.set_text_color(0, 0, 0)


def _style_heading_color(paragraph, rgb_tuple) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*rgb_tuple)


def build_docx(cover_blocks, body_blocks):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    if cover_blocks:
        info = _extract_cover_info(cover_blocks)
        _build_docx_cover(doc, info)

    for kind, payload in body_blocks:
        if kind == "h1":
            p = doc.add_heading(level=0)
            render_runs_docx(p, payload)
            _style_heading_color(p, PALETTE["h1_rgb"])
        elif kind == "h2":
            p = doc.add_heading(level=1)
            render_runs_docx(p, payload)
            _style_heading_color(p, PALETTE["h2_rgb"])
        elif kind == "h3":
            p = doc.add_heading(level=2)
            render_runs_docx(p, payload)
            _style_heading_color(p, PALETTE["h3_rgb"])
        elif kind == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            render_runs_docx(p, payload)
        elif kind == "ul":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                render_runs_docx(p, item)
        elif kind == "ol":
            for item in payload:
                p = doc.add_paragraph(style="List Number")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                render_runs_docx(p, item)
        elif kind == "quote":
            p = doc.add_paragraph(style="Intense Quote")
            render_runs_docx(p, payload)
        elif kind == "code":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif kind == "math":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.italic = True
            r.font.name = "Cambria Math"
        elif kind == "hr":
            p = doc.add_paragraph("_" * 70)
            p.alignment = 1  # center
        elif kind == "table":
            if not payload:
                continue
            _render_docx_table(doc, payload)
            doc.add_paragraph()  # spacer after table

    doc.save(DOCX)
    print(f"wrote {DOCX} ({DOCX.stat().st_size/1024:.1f} KB)")


# ---------- PDF rendering ------------------------------------------------------

class ReportPDF(FPDF):
    def __init__(self):
        super().__init__(format="A4")
        self.set_margins(20, 18, 20)
        self.set_auto_page_break(True, margin=18)
        # Embed Arial so we get full unicode
        self.add_font("Arial", "", ARIAL_REG, uni=True)
        self.add_font("Arial", "B", ARIAL_BOLD, uni=True)
        self.add_font("Arial", "I", ARIAL_IT, uni=True)
        self.add_font("Arial", "BI", ARIAL_BI, uni=True)
        self.add_font("Mono", "", CONSOLAS, uni=True)
        self.set_font("Arial", "", 11)

    def add_inline(self, text: str, base_style: str = ""):
        """Render text with inline **/*/` formatting on the current line.

        Wraps automatically when text exceeds the line width.
        """
        for kind, content in inline_tokens(text):
            if kind == "link":
                content = content[0]
                style = base_style
                self.set_text_color(*PALETTE["link_rgb"])
                self._write_wrapped(content, "Arial", style, 11)
                self.set_text_color(0, 0, 0)
            elif kind == "bold":
                self._write_wrapped(content, "Arial", "B", 11)
            elif kind == "italic":
                self._write_wrapped(content, "Arial", "I", 11)
            elif kind == "code":
                self._write_wrapped(content, "Mono", "", 10)
            else:
                self._write_wrapped(content, "Arial", base_style, 11)

    def _write_wrapped(self, text: str, family: str, style: str, size: int):
        self.set_font(family, style, size)
        # write() auto-wraps; line height 5.2mm
        self.write(5.2, text)


# ---------- Justified inline writer ------------------------------------------
# Word-wrapped paragraph renderer that keeps **bold**, *italic*, `code`, and
# [link](url) formatting while distributing extra space across word gaps so
# each non-final line is fully justified.

def _font_for(kind: str):
    if kind in ("bold", "link"):
        return ("Arial", "B", 11)
    if kind == "italic":
        return ("Arial", "I", 11)
    if kind == "code":
        return ("Mono", "", 10)
    return ("Arial", "", 11)


def _color_for(kind: str):
    if kind == "bold":
        return PALETTE["h1_rgb"]
    if kind == "code":
        return (180, 83, 9)  # amber
    if kind == "link":
        return PALETTE["link_rgb"]
    return (0, 0, 0)


def _tokenize_atoms(text: str):
    """Break text into (kind, word_or_space) atoms preserving inline style."""
    atoms = []
    for kind, content in inline_tokens(text):
        if kind == "link":
            content = content[0]
        for part in re.split(r"(\s+)", content):
            if part == "":
                continue
            atoms.append((kind, part))
    return atoms


def write_justified(pdf, text: str, line_h: float = 5.4) -> None:
    avail = pdf.w - pdf.l_margin - pdf.r_margin
    atoms = _tokenize_atoms(text)
    # Measure each atom.
    measured = []
    for kind, word in atoms:
        fam, sty, sz = _font_for(kind)
        pdf.set_font(fam, sty, sz)
        w = pdf.get_string_width(word)
        measured.append((kind, word, w, word.isspace()))

    # Pack atoms into lines that fit within `avail`.
    lines, cur, cur_w = [], [], 0.0
    for atom in measured:
        _, _, w, is_space = atom
        if cur_w + w > avail and cur:
            while cur and cur[-1][3]:
                cur.pop()
            lines.append(cur)
            cur, cur_w = [], 0.0
            if is_space:
                continue
        cur.append(atom)
        cur_w += w
    if cur:
        while cur and cur[-1][3]:
            cur.pop()
        lines.append(cur)

    # Render each line, justifying all but the last. Treat each line as atomic:
    # if it doesn't fit on the current page, force a page break BEFORE writing
    # the line, so the captured y can't go stale mid-line.
    bottom_trigger = pdf.h - pdf.b_margin
    for idx, line in enumerate(lines):
        if pdf.get_y() + line_h > bottom_trigger:
            pdf.add_page()

        is_last = idx == len(lines) - 1
        non_space_w = sum(w for _, _, w, sp in line if not sp)
        n_spaces = sum(1 for _, _, _, sp in line if sp)
        base_space_w = sum(w for _, _, w, sp in line if sp)
        if is_last or n_spaces == 0:
            extra_per_space = 0.0
        else:
            extra_per_space = (avail - non_space_w - base_space_w) / n_spaces

        x = pdf.l_margin
        y = pdf.get_y()
        for kind, word, w, is_space in line:
            fam, sty, sz = _font_for(kind)
            pdf.set_font(fam, sty, sz)
            pdf.set_text_color(*_color_for(kind))
            if is_space:
                x += w + extra_per_space
            else:
                pdf.set_xy(x, y)
                pdf.cell(w + 0.1, line_h, word)
                x += w
        pdf.set_text_color(0, 0, 0)
        pdf.set_xy(pdf.l_margin, y + line_h)


def render_pdf(cover_blocks, body_blocks):
    pdf = ReportPDF()
    pdf.add_page()

    if cover_blocks:
        info = _extract_cover_info(cover_blocks)
        _render_pdf_cover(pdf, info)
        pdf.add_page()

    for kind, payload in body_blocks:
        if kind == "h1":
            pdf.set_font("Arial", "B", 22)
            pdf.set_text_color(*PALETTE["h1_rgb"])
            pdf.multi_cell(0, 11, payload)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
        elif kind == "h2":
            pdf.ln(3)
            # Atomic-fit so the heading and its underline never land on
            # different pages.
            pdf.set_font("Arial", "B", 15)
            avail_w = pdf.w - pdf.l_margin - pdf.r_margin
            lines = pdf.multi_cell(avail_w, 8, payload,
                                   dry_run=True, output="LINES")
            needed = max(8, len(lines) * 8) + 6
            if pdf.get_y() + needed > pdf.h - pdf.b_margin:
                pdf.add_page()
            pdf.set_font("Arial", "B", 15)
            pdf.set_text_color(*PALETTE["h2_rgb"])
            pdf.multi_cell(0, 8, payload)
            # Decorative amber underline beneath each h2
            y = pdf.get_y() + 0.5
            r, g, b = PALETTE["cover_band_b_rgb"]
            pdf.set_draw_color(r, g, b)
            pdf.set_line_width(0.6)
            pdf.line(pdf.l_margin, y, pdf.l_margin + 40, y)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(3)
        elif kind == "h3":
            pdf.ln(2)
            pdf.set_font("Arial", "B", 12)
            pdf.set_text_color(*PALETTE["h3_rgb"])
            pdf.multi_cell(0, 6, payload)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(0.5)
        elif kind == "p":
            write_justified(pdf, payload)
            pdf.ln(2)
        elif kind == "ul":
            for item in payload:
                pdf.set_font("Arial", "", 11)
                pdf.cell(5, 5.2, "")
                r, g, b = PALETTE["h3_rgb"]
                pdf.set_text_color(r, g, b)
                pdf.write(5.2, "•  ")
                pdf.set_text_color(0, 0, 0)
                pdf.add_inline(item)
                pdf.ln(5.5)
            pdf.ln(2)
        elif kind == "ol":
            for n, item in enumerate(payload, 1):
                pdf.set_font("Arial", "", 11)
                pdf.cell(5, 5.2, "")
                r, g, b = PALETTE["h3_rgb"]
                pdf.set_text_color(r, g, b)
                pdf.write(5.2, f"{n}.  ")
                pdf.set_text_color(0, 0, 0)
                pdf.add_inline(item)
                pdf.ln(5.5)
            pdf.ln(2)
        elif kind == "quote":
            pdf.set_font("Arial", "I", 11)
            pdf.set_text_color(*PALETTE["quote_rgb"])
            pdf.set_left_margin(25)
            # Left teal bar before the quote
            y0 = pdf.get_y()
            r, g, b = PALETTE["h2_rgb"]
            pdf.set_draw_color(r, g, b)
            pdf.set_line_width(0.8)
            pdf.line(22, y0, 22, y0 + 5)
            pdf.multi_cell(0, 5.5, payload)
            pdf.set_left_margin(20)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(3)
        elif kind == "code":
            r, g, b = PALETTE["code_fill_rgb"]
            pdf.set_fill_color(r, g, b)
            pdf.set_font("Mono", "", 9)
            for code_line in payload.split("\n"):
                pdf.cell(0, 4.5, code_line, ln=1, fill=True)
            pdf.ln(3)
        elif kind == "math":
            pdf.set_font("Arial", "I", 11)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(0, 5.5, payload)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(3)
        elif kind == "hr":
            y = pdf.get_y() + 2
            r, g, b = PALETTE["hr_rgb"]
            pdf.set_draw_color(r, g, b)
            pdf.set_line_width(0.6)
            pdf.line(20, y, 190, y)
            pdf.ln(6)
        elif kind == "table":
            render_table_pdf(pdf, payload)

    pdf.output(str(PDF))
    print(f"wrote {PDF} ({PDF.stat().st_size/1024:.1f} KB)")


def _strip_inline(text: str) -> str:
    """Strip markdown for a plain-text snapshot used in table cells."""
    parts = []
    for kind, content in inline_tokens(text):
        parts.append(content[0] if kind == "link" else content)
    return "".join(parts)


def _compute_col_widths(pdf, rows, n_cols, page_width):
    """Allocate column widths in proportion to widest cell content, with floors."""
    pdf.set_font("Arial", "", 10)
    raw = [0.0] * n_cols
    for row in rows:
        for c_idx in range(n_cols):
            cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
            for line in cell.splitlines() or [cell]:
                w = pdf.get_string_width(line)
                if w > raw[c_idx]:
                    raw[c_idx] = w

    pad = 4.0  # mm of horizontal padding per cell
    raw = [w + pad for w in raw]
    total = sum(raw) or 1.0
    # Floor at 14mm so even tiny numeric columns are readable.
    widths = [max(14.0, w * page_width / total) for w in raw]
    # Re-normalize to exactly fill the page.
    scale = page_width / sum(widths)
    return [w * scale for w in widths]


def render_table_pdf(pdf, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    widths = _compute_col_widths(pdf, rows, n_cols, page_width)

    # All cells centered in both header and body.
    col_align = ["C"] * n_cols

    pdf.ln(2)
    line_h = 5.4
    v_pad = 1.6  # extra vertical padding inside a cell

    def measure(row, font_style):
        pdf.set_font("Arial", font_style, 10)
        max_lines = 1
        for c_idx in range(n_cols):
            cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
            # multi_cell wrap measurement
            lines = pdf.multi_cell(widths[c_idx], line_h, cell, dry_run=True, output="LINES")
            if len(lines) > max_lines:
                max_lines = len(lines)
        return max_lines * line_h + v_pad

    for r_idx, row in enumerate(rows):
        is_header = r_idx == 0
        style = "B" if is_header else ""
        h = measure(row, style)
        if pdf.get_y() + h > pdf.h - pdf.b_margin:
            pdf.add_page()
            # Re-emit the header on the new page so the table stays readable.
            if not is_header and rows:
                hh = measure(rows[0], "B")
                _draw_pdf_row(pdf, rows[0], widths, col_align, hh, line_h, v_pad,
                              is_header=True)

        _draw_pdf_row(pdf, row, widths, col_align, h, line_h, v_pad, is_header=is_header,
                      zebra=(r_idx % 2 == 0 and not is_header))
    pdf.ln(4)


def _draw_pdf_row(pdf, row, widths, col_align, h, line_h, v_pad, is_header=False, zebra=False):
    n_cols = len(widths)
    y_start = pdf.get_y()
    for c_idx in range(n_cols):
        cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
        x_start = pdf.l_margin + sum(widths[:c_idx])
        if is_header:
            r, g, b = PALETTE["tbl_header_rgb"]
            pdf.set_fill_color(r, g, b)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Arial", "B", 10)
        else:
            if zebra:
                r, g, b = PALETTE["tbl_zebra_rgb"]
                pdf.set_fill_color(r, g, b)
            else:
                pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(20, 20, 20)
            pdf.set_font("Arial", "", 10)
        r, g, b = PALETTE["tbl_border_rgb"]
        pdf.set_draw_color(r, g, b)
        pdf.set_line_width(0.2)

        align = "C"

        # Draw the filled border first so the text sits cleanly inside.
        pdf.rect(x_start, y_start, widths[c_idx], h, "DF")

        # Wrap text inside the cell with a small inset so it doesn't touch the border.
        inset_x = 1.2
        inset_y = v_pad / 2
        pdf.set_xy(x_start + inset_x, y_start + inset_y)
        pdf.multi_cell(widths[c_idx] - 2 * inset_x, line_h, cell, border=0, align=align,
                       fill=False)
    pdf.set_xy(pdf.l_margin, y_start + h)
    pdf.set_text_color(0, 0, 0)


# ---------- Main ---------------------------------------------------------------

EMOJI_FALLBACKS = {
    "✅": "[OK]",
    "❗": "[!]",
    "\U0001f150": "(A)",  # 🅰
    "\U0001f151": "(B)",  # 🅱
    "\U0001f7e2": "(A)",  # 🟢
    "\U0001f7e1": "(B)",  # 🟡
    "\U0001f534": "(C)",  # 🔴
    "\U0001f4c1": "[folder]",  # 📁
    "⚠️": "[!]",  # warning
    "⚠": "[!]",
}


def normalize_emoji(text: str) -> str:
    for k, v in EMOJI_FALLBACKS.items():
        text = text.replace(k, v)
    return text


def main():
    md_text = MD.read_text(encoding="utf-8")
    md_text = normalize_emoji(md_text)
    blocks = list(parse(md_text))
    cover_blocks, body_blocks = _split_cover(blocks)
    print(f"Parsed {len(blocks)} blocks "
          f"(cover={len(cover_blocks)}, body={len(body_blocks)}) from {MD}")

    build_docx(cover_blocks, body_blocks)
    render_pdf(cover_blocks, body_blocks)


if __name__ == "__main__":
    main()
