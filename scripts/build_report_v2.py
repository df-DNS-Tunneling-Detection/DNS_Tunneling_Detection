"""Premium-styled DOCX + PDF build of reports/report.md.

Reads the same markdown as build_report.py but renders a fancier cover page,
running page header/footer, pilled section headings, callout-style block quotes,
and a more polished table treatment. Writes to:

    reports/report_v2.docx
    reports/report_v2.pdf

The original report.docx / report.pdf are NOT touched.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

# Reuse the markdown parser, inline tokenizer, cover-extraction, and the
# OxmlElement helpers from the v1 builder. Only the *rendering* is new here.
sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_report as v1  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402
from fpdf import FPDF  # noqa: E402


ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "reports" / "report.md"
DOCX = ROOT / "reports" / "report_v2.docx"
PDF = ROOT / "reports" / "report_v2.pdf"


# ---------- Premium palette ---------------------------------------------------
# Indigo primary + rose accent + warm slate text. Distinct from v1's teal scheme.

P = {
    # Primaries
    "indigo_900":     (49, 46, 129),   # #312E81
    "indigo_700":     (67, 56, 202),   # #4338CA
    "indigo_500":     (99, 102, 241),  # #6366F1
    "indigo_100":     (224, 231, 255), # #E0E7FF
    "indigo_50":      (238, 242, 255), # #EEF2FF
    # Accent
    "rose_600":       (225, 29, 72),   # #E11D48
    "rose_400":       (251, 113, 133), # #FB7185
    "amber_500":      (245, 158, 11),  # #F59E0B
    # Neutrals
    "slate_900":      (15, 23, 42),    # #0F172A
    "slate_700":      (51, 65, 85),    # #334155
    "slate_500":      (100, 116, 139), # #64748B
    "slate_200":      (226, 232, 240), # #E2E8F0
    "slate_100":      (241, 245, 249), # #F1F5F9
    # Hex forms used for DOCX XML
    "indigo_900_hex": "312E81",
    "indigo_700_hex": "4338CA",
    "indigo_100_hex": "E0E7FF",
    "indigo_50_hex":  "EEF2FF",
    "rose_600_hex":   "E11D48",
    "amber_500_hex":  "F59E0B",
    "slate_200_hex":  "E2E8F0",
}

# Footer / header text shown on every body page
RUNNING_HEADER = "DNS Tunneling Detection · Project Report"
RUNNING_FOOTER = "Digital Forensics — 2026"


# ---------- Shared helpers from v1 -------------------------------------------

parse = v1.parse
inline_tokens = v1.inline_tokens
_strip_inline = v1._strip_inline
_is_numeric_cell = v1._is_numeric_cell
_split_cover = v1._split_cover
_extract_cover_info = v1._extract_cover_info
normalize_emoji = v1.normalize_emoji


# ---------- v2-flavoured justified writer ------------------------------------
# Reuses v1.write_justified semantics but swaps the font/color tables for the
# premium indigo+rose palette.

def _v2_font_for(kind: str):
    if kind in ("bold", "link"):
        return ("Arial", "B", 11)
    if kind == "italic":
        return ("Arial", "I", 11)
    if kind == "code":
        return ("Mono", "", 10)
    return ("Arial", "", 11)


def _v2_color_for(kind: str):
    if kind == "bold":
        return P["indigo_900"]
    if kind == "code":
        return P["rose_600"]
    if kind == "link":
        return P["indigo_700"]
    return P["slate_900"]


def write_justified_v2(pdf, text: str, line_h: float = 5.4) -> None:
    avail = pdf.w - pdf.l_margin - pdf.r_margin

    atoms = []
    for kind, content in inline_tokens(text):
        if kind == "link":
            content = content[0]
        for part in re.split(r"(\s+)", content):
            if part == "":
                continue
            atoms.append((kind, part))

    measured = []
    for kind, word in atoms:
        fam, sty, sz = _v2_font_for(kind)
        pdf.set_font(fam, sty, sz)
        w = pdf.get_string_width(word)
        measured.append((kind, word, w, word.isspace()))

    lines, cur, cur_w = [], [], 0.0
    for atom in measured:
        _, _, w, is_space = atom
        if cur_w + w > avail and cur:
            while cur and cur[-1][3]:
                cur.pop()
            lines.append(cur)
            cur, cur_w = [], 0.0
            if is_space:
                continue
        cur.append(atom)
        cur_w += w
    if cur:
        while cur and cur[-1][3]:
            cur.pop()
        lines.append(cur)

    bottom_trigger = pdf.h - pdf.b_margin
    for idx, line in enumerate(lines):
        if pdf.get_y() + line_h > bottom_trigger:
            pdf.add_page()

        is_last = idx == len(lines) - 1
        non_space_w = sum(w for _, _, w, sp in line if not sp)
        n_spaces = sum(1 for _, _, _, sp in line if sp)
        base_space_w = sum(w for _, _, w, sp in line if sp)
        if is_last or n_spaces == 0:
            extra_per_space = 0.0
        else:
            extra_per_space = (avail - non_space_w - base_space_w) / n_spaces

        x = pdf.l_margin
        y = pdf.get_y()
        for kind, word, w, is_space in line:
            fam, sty, sz = _v2_font_for(kind)
            pdf.set_font(fam, sty, sz)
            pdf.set_text_color(*_v2_color_for(kind))
            if is_space:
                x += w + extra_per_space
            else:
                pdf.set_xy(x, y)
                pdf.cell(w + 0.1, line_h, word)
                x += w
        pdf.set_text_color(0, 0, 0)
        pdf.set_xy(pdf.l_margin, y + line_h)


# ---------- DOCX low-level helpers (mirrors v1 but kept local) ---------------

def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=120, bottom=120, left=160, right=160) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _set_table_borders(table, color_hex: str, size: str = "4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color_hex)
        borders.append(node)
    tbl_pr.append(borders)


def _add_paragraph_shading(paragraph, hex_fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    p_pr.append(shd)


def _add_centered(doc, text, *, size=12, bold=False, italic=False,
                  rgb=(0, 0, 0), space_after=4, space_before=0,
                  font_name="Calibri"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*rgb)
    return p


def _render_runs_docx(paragraph, text: str, *, base_color=None) -> None:
    """Inline markdown rendering for DOCX paragraphs."""
    for kind, content in inline_tokens(text):
        if kind == "link":
            label, _url = content
            r = paragraph.add_run(label)
            r.font.color.rgb = RGBColor(*P["indigo_700"])
            r.underline = True
        elif kind == "bold":
            r = paragraph.add_run(content)
            r.bold = True
            if base_color is not None:
                r.font.color.rgb = RGBColor(*base_color)
        elif kind == "italic":
            r = paragraph.add_run(content)
            r.italic = True
            if base_color is not None:
                r.font.color.rgb = RGBColor(*base_color)
        elif kind == "code":
            r = paragraph.add_run(content)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(*P["rose_600"])
        else:
            r = paragraph.add_run(content)
            if base_color is not None:
                r.font.color.rgb = RGBColor(*base_color)


# ---------- DOCX cover page --------------------------------------------------

def _build_docx_cover(doc, info: dict) -> None:
    # Top double band: indigo_900 + rose accent stripe
    band = doc.add_paragraph()
    band.paragraph_format.space_before = Pt(0)
    band.paragraph_format.space_after = Pt(0)
    r = band.add_run(" " * 6)
    r.font.size = Pt(16)
    _add_paragraph_shading(band, P["indigo_900_hex"])

    stripe = doc.add_paragraph()
    stripe.paragraph_format.space_before = Pt(0)
    stripe.paragraph_format.space_after = Pt(28)
    r = stripe.add_run(" " * 6)
    r.font.size = Pt(4)
    _add_paragraph_shading(stripe, P["rose_600_hex"])

    # Vertical padding before title
    for _ in range(2):
        doc.add_paragraph()

    # Title — two-tone (indigo primary, rose accent for the colon clause)
    title = info["title"]
    if ":" in title:
        primary, _, accent = title.partition(":")
        accent = accent.strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(primary.strip())
        r1.font.name = "Calibri"
        r1.font.size = Pt(30)
        r1.bold = True
        r1.font.color.rgb = RGBColor(*P["indigo_900"])
        _add_centered(doc, accent, size=22, italic=True,
                      rgb=P["rose_600"], space_after=10)
    else:
        _add_centered(doc, title, size=30, bold=True,
                      rgb=P["indigo_900"], space_after=10)

    # Decorative double rule
    _add_centered(doc, "━" * 18, size=12,
                  rgb=P["indigo_700"], space_after=0)
    _add_centered(doc, "━" * 9, size=10,
                  rgb=P["rose_600"], space_after=22)

    # Course banner
    _add_centered(doc, info.get("course", ""), size=17, italic=True,
                  rgb=P["indigo_700"], space_after=4)
    _add_centered(doc, "Project Report", size=12, italic=True,
                  rgb=P["slate_500"], space_after=30)

    # Supervisor — boxed card via 1x1 table with shading + border
    sup_tbl = doc.add_table(rows=2, cols=1)
    sup_tbl.alignment = 1  # WD_TABLE_ALIGNMENT.CENTER
    sup_tbl.autofit = False
    sup_tbl.columns[0].width = Inches(3.6)

    c0 = sup_tbl.cell(0, 0)
    c0.width = Inches(3.6)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c0.text = ""
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run("SUBMITTED TO")
    rr.bold = True
    rr.font.name = "Calibri"
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(*P["rose_600"])
    _set_cell_shading(c0, P["indigo_50_hex"])
    _set_cell_margins(c0, top=120, bottom=40, left=160, right=160)

    c1 = sup_tbl.cell(1, 0)
    c1.width = Inches(3.6)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.text = ""
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(info.get("supervisor", ""))
    rr.bold = True
    rr.font.name = "Calibri"
    rr.font.size = Pt(16)
    rr.font.color.rgb = RGBColor(*P["indigo_900"])
    _set_cell_shading(c1, P["indigo_50_hex"])
    _set_cell_margins(c1, top=20, bottom=160, left=160, right=160)

    _set_table_borders(sup_tbl, P["indigo_700_hex"], size="6")

    # spacer
    doc.add_paragraph()

    # Date — small centered line
    _add_centered(doc, "DATE", size=10, bold=True,
                  rgb=P["rose_600"], space_after=2)
    _add_centered(doc, info.get("date", ""), size=14,
                  rgb=P["slate_700"], space_after=24)

    # Team section
    _add_centered(doc, "PROJECT TEAM", size=12, bold=True,
                  rgb=P["rose_600"], space_after=4)
    _add_centered(doc, "━" * 10, size=10,
                  rgb=P["indigo_700"], space_after=8)

    # Team list: 2-column centered table so badges and names line up vertically.
    team_tbl = doc.add_table(rows=len(info["team"]), cols=2)
    team_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    team_tbl.autofit = False
    team_tbl.columns[0].width = Inches(0.45)
    team_tbl.columns[1].width = Inches(2.8)

    for r_idx, name in enumerate(info["team"]):
        c0 = team_tbl.cell(r_idx, 0)
        c0.width = Inches(0.45)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c0.text = ""
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        rn = p.add_run(f"{r_idx + 1:02d}")
        rn.bold = True
        rn.font.name = "Calibri"
        rn.font.size = Pt(11)
        rn.font.color.rgb = RGBColor(*P["rose_600"])
        _set_cell_margins(c0, top=20, bottom=20, left=80, right=80)

        c1 = team_tbl.cell(r_idx, 1)
        c1.width = Inches(2.8)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.text = ""
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        rname = p.add_run(name)
        rname.font.name = "Calibri"
        rname.font.size = Pt(13)
        rname.font.color.rgb = RGBColor(*P["slate_900"])
        _set_cell_margins(c1, top=20, bottom=20, left=140, right=80)

    # Hide table borders.
    tbl_el = team_tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)

    # Push to bottom — a few empties then the footer band
    for _ in range(3):
        doc.add_paragraph()

    # Bottom rose stripe + indigo band with repository URL
    stripe = doc.add_paragraph()
    stripe.paragraph_format.space_before = Pt(0)
    stripe.paragraph_format.space_after = Pt(0)
    r = stripe.add_run(" " * 6)
    r.font.size = Pt(4)
    _add_paragraph_shading(stripe, P["rose_600_hex"])

    bottom = doc.add_paragraph()
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom.paragraph_format.space_before = Pt(0)
    bottom.paragraph_format.space_after = Pt(0)
    r = bottom.add_run("  " + _clean_url(info.get("repository", "")) + "  ")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    _add_paragraph_shading(bottom, P["indigo_900_hex"])

    # Page break
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------- DOCX body --------------------------------------------------------

def _render_docx_table_premium(doc, rows) -> None:
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = tbl.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx == 0:
                _render_runs_docx(p, cell_text, base_color=(255, 255, 255))
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(11)
                _set_cell_shading(cell, P["indigo_700_hex"])
            else:
                _render_runs_docx(p, cell_text, base_color=P["slate_900"])
                for run in p.runs:
                    run.font.size = Pt(10.5)
                if r_idx % 2 == 0:
                    _set_cell_shading(cell, P["indigo_50_hex"])
            _set_cell_margins(cell, top=110, bottom=110, left=180, right=180)

    _set_table_borders(tbl, P["indigo_100_hex"], size="6")


def _render_docx_quote(doc, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(0.12)
    tbl.columns[1].width = Inches(5.8)

    bar = tbl.cell(0, 0)
    bar.width = Inches(0.12)
    bar.text = ""
    _set_cell_shading(bar, P["rose_600_hex"])
    _set_cell_margins(bar, top=0, bottom=0, left=0, right=0)

    body = tbl.cell(0, 1)
    body.width = Inches(5.8)
    body.text = ""
    p = body.paragraphs[0]
    _render_runs_docx(p, text, base_color=P["slate_700"])
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(11)
    _set_cell_shading(body, P["indigo_50_hex"])
    _set_cell_margins(body, top=160, bottom=160, left=200, right=200)

    # No borders on the callout
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def _style_heading_color(paragraph, rgb_tuple) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*rgb_tuple)


def build_docx(cover_blocks, body_blocks):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)

    if cover_blocks:
        info = _extract_cover_info(cover_blocks)
        _build_docx_cover(doc, info)

    for kind, payload in body_blocks:
        if kind == "h1":
            p = doc.add_heading(level=0)
            _render_runs_docx(p, payload, base_color=P["indigo_900"])
            _style_heading_color(p, P["indigo_900"])
        elif kind == "h2":
            p = doc.add_heading(level=1)
            _render_runs_docx(p, payload, base_color=P["indigo_700"])
            _style_heading_color(p, P["indigo_700"])
        elif kind == "h3":
            p = doc.add_heading(level=2)
            _render_runs_docx(p, payload, base_color=P["rose_600"])
            _style_heading_color(p, P["rose_600"])
        elif kind == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _render_runs_docx(p, payload, base_color=P["slate_900"])
        elif kind == "ul":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _render_runs_docx(p, item, base_color=P["slate_900"])
        elif kind == "ol":
            for item in payload:
                p = doc.add_paragraph(style="List Number")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _render_runs_docx(p, item, base_color=P["slate_900"])
        elif kind == "quote":
            _render_docx_quote(doc, payload)
        elif kind == "code":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(*P["slate_700"])
            _add_paragraph_shading(p, P["slate_100"][:1] and "F1F5F9")
        elif kind == "math":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.italic = True
            r.font.name = "Cambria Math"
            r.font.color.rgb = RGBColor(*P["slate_700"])
        elif kind == "hr":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("◆  ◆  ◆")
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(*P["rose_600"])
        elif kind == "table":
            if not payload:
                continue
            _render_docx_table_premium(doc, payload)
            doc.add_paragraph()

    doc.save(DOCX)
    print(f"wrote {DOCX} ({DOCX.stat().st_size/1024:.1f} KB)")


# ---------- PDF --------------------------------------------------------------

class PremiumPDF(FPDF):
    def __init__(self):
        super().__init__(format="A4")
        self.set_margins(22, 24, 22)
        self.set_auto_page_break(True, margin=22)
        # Embed Arial for unicode coverage.
        self.add_font("Arial", "", v1.ARIAL_REG, uni=True)
        self.add_font("Arial", "B", v1.ARIAL_BOLD, uni=True)
        self.add_font("Arial", "I", v1.ARIAL_IT, uni=True)
        self.add_font("Arial", "BI", v1.ARIAL_BI, uni=True)
        self.add_font("Mono", "", v1.CONSOLAS, uni=True)
        self.set_font("Arial", "", 11)
        # Cover page suppresses header/footer; toggled via this flag.
        self._chrome = False

    def header(self):
        # Skip on cover (page 1) — FPDF calls header() of the next page
        # during add_page(), so we can't reliably toggle _chrome around it.
        if self.page_no() <= 1 or not self._chrome:
            return
        # Thin indigo line + small caps title on the left, page number right
        x0, x1 = self.l_margin, self.w - self.r_margin
        y = 14
        r, g, b = P["indigo_700"]
        self.set_draw_color(r, g, b)
        self.set_line_width(0.5)
        self.line(x0, y, x1, y)

        self.set_xy(x0, y - 7)
        self.set_font("Arial", "B", 8)
        self.set_text_color(*P["indigo_900"])
        self.cell((x1 - x0) / 2, 5, RUNNING_HEADER.upper(), align="L")

        self.set_xy((x0 + x1) / 2, y - 7)
        self.set_font("Arial", "", 8)
        self.set_text_color(*P["slate_500"])
        self.cell((x1 - x0) / 2, 5, f"Page {self.page_no() - 1}", align="R")
        self.set_text_color(0, 0, 0)
        self.set_y(self.t_margin)

    def footer(self):
        if self.page_no() <= 1 or not self._chrome:
            return
        x0, x1 = self.l_margin, self.w - self.r_margin
        y = self.h - 14
        r, g, b = P["rose_600"]
        self.set_draw_color(r, g, b)
        self.set_line_width(0.4)
        self.line(x0, y, x1, y)

        self.set_xy(x0, y + 2)
        self.set_font("Arial", "I", 8)
        self.set_text_color(*P["slate_500"])
        self.cell(x1 - x0, 4, RUNNING_FOOTER, align="C")
        self.set_text_color(0, 0, 0)

    # --- inline rendering -----------------------------------------------------

    def _line_h(self):
        return 5.4

    def add_inline(self, text: str, base_style: str = ""):
        for kind, content in inline_tokens(text):
            if kind == "link":
                label = content[0]
                self.set_text_color(*P["indigo_700"])
                self._write(label, "Arial", base_style or "B", 11)
                self.set_text_color(0, 0, 0)
            elif kind == "bold":
                self.set_text_color(*P["indigo_900"])
                self._write(content, "Arial", "B", 11)
                self.set_text_color(0, 0, 0)
            elif kind == "italic":
                self._write(content, "Arial", "I", 11)
            elif kind == "code":
                self.set_text_color(*P["rose_600"])
                self._write(content, "Mono", "", 10)
                self.set_text_color(0, 0, 0)
            else:
                self._write(content, "Arial", base_style, 11)

    def _write(self, text: str, family: str, style: str, size: int):
        self.set_font(family, style, size)
        self.write(self._line_h(), text)


# ---------- PDF cover --------------------------------------------------------

def _render_pdf_cover(pdf: PremiumPDF, info: dict) -> None:
    # Disable auto page-break + zero side margins so the cover's full-width
    # `cell(page_w, …, align="C")` calls actually center on the page.
    saved_apb = pdf.auto_page_break
    saved_b = pdf.b_margin
    saved_l = pdf.l_margin
    saved_r = pdf.r_margin
    pdf.set_auto_page_break(False)
    pdf.set_left_margin(0)
    pdf.set_right_margin(0)
    try:
        _render_pdf_cover_impl(pdf, info)
    finally:
        pdf.set_left_margin(saved_l)
        pdf.set_right_margin(saved_r)
        pdf.set_auto_page_break(saved_apb, margin=saved_b)


def _clean_url(url: str) -> str:
    return (url or "").strip().strip("<>").strip()


def _render_pdf_cover_impl(pdf: PremiumPDF, info: dict) -> None:
    page_w, page_h = pdf.w, pdf.h
    cx = page_w / 2

    # Top decorative double band
    r, g, b = P["indigo_900"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, 0, page_w, 28.0, "F")

    r, g, b = P["rose_600"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, 28.0, page_w, 4.0, "F")

    # Decorative geometric corner accent (top-right) — three small rose squares
    r, g, b = P["rose_400"]
    pdf.set_fill_color(r, g, b)
    for i in range(3):
        pdf.rect(page_w - 14 - i * 7, 10 - i * 2, 4, 4, "F")

    # Title — two-tone if it contains a colon
    title = info["title"]
    if ":" in title:
        primary, _, accent = title.partition(":")
        accent = accent.strip()
        pdf.set_y(70.0)
        pdf.set_font("Arial", "B", 28)
        pdf.set_text_color(*P["indigo_900"])
        pdf.cell(page_w, 12, primary.strip(), align="C", ln=1)
        pdf.set_font("Arial", "BI", 20)
        pdf.set_text_color(*P["rose_600"])
        pdf.cell(page_w, 10, accent, align="C", ln=1)
    else:
        pdf.set_y(74.0)
        pdf.set_font("Arial", "B", 28)
        pdf.set_text_color(*P["indigo_900"])
        pdf.cell(page_w, 12, title, align="C", ln=1)

    # Decorative double rule under title
    y = pdf.get_y() + 4
    r, g, b = P["indigo_700"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(1.0)
    pdf.line(cx - 40, y, cx + 40, y)
    r, g, b = P["rose_600"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.6)
    pdf.line(cx - 20, y + 2.5, cx + 20, y + 2.5)
    pdf.set_y(y + 10)

    # Course banner
    pdf.set_font("Arial", "I", 17)
    pdf.set_text_color(*P["indigo_700"])
    pdf.cell(page_w, 9, info.get("course", ""), align="C", ln=1)
    pdf.set_font("Arial", "I", 11)
    pdf.set_text_color(*P["slate_500"])
    pdf.cell(page_w, 6, "Project Report", align="C", ln=1)
    pdf.ln(18)

    # Supervisor card (boxed) — indigo_50 fill with indigo_700 border
    card_w = 130.0
    card_h = 30.0
    card_x = (page_w - card_w) / 2
    card_y = pdf.get_y()
    r, g, b = P["indigo_50"]
    pdf.set_fill_color(r, g, b)
    r, g, b = P["indigo_700"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.6)
    pdf.rect(card_x, card_y, card_w, card_h, "DF")

    pdf.set_xy(card_x, card_y + 5)
    pdf.set_font("Arial", "B", 9)
    pdf.set_text_color(*P["rose_600"])
    pdf.cell(card_w, 5, "SUBMITTED TO", align="C")

    pdf.set_xy(card_x, card_y + 13)
    pdf.set_font("Arial", "B", 17)
    pdf.set_text_color(*P["indigo_900"])
    pdf.cell(card_w, 10, info.get("supervisor", ""), align="C")

    pdf.set_y(card_y + card_h + 8)

    # Date
    pdf.set_font("Arial", "B", 9)
    pdf.set_text_color(*P["rose_600"])
    pdf.cell(page_w, 5, "DATE", align="C", ln=1)
    pdf.set_font("Arial", "", 13)
    pdf.set_text_color(*P["slate_700"])
    pdf.cell(page_w, 7, info.get("date", ""), align="C", ln=1)
    pdf.ln(8)

    # Team header
    pdf.set_font("Arial", "B", 11)
    pdf.set_text_color(*P["rose_600"])
    pdf.cell(page_w, 6, "PROJECT TEAM", align="C", ln=1)

    y_line = pdf.get_y() + 1
    r, g, b = P["indigo_700"]
    pdf.set_draw_color(r, g, b)
    pdf.set_line_width(0.5)
    pdf.line(cx - 18, y_line, cx + 18, y_line)
    pdf.set_y(y_line + 4)

    # Each team member: numeric badge + name, every row using the same
    # x_start so badges and names align in a clean column centered on the page.
    badge_d = 6.0
    gap = 4.0
    row_h = 8.0
    pdf.set_font("Arial", "", 13)
    max_name_w = max(pdf.get_string_width(n) for n in info["team"])
    block_w = badge_d + gap + max_name_w
    x_block = (page_w - block_w) / 2

    for n, name in enumerate(info["team"], 1):
        y_row = pdf.get_y()

        # Filled rose circle as the badge (fixed x)
        r, g, b = P["rose_600"]
        pdf.set_fill_color(r, g, b)
        pdf.set_draw_color(r, g, b)
        pdf.ellipse(x_block, y_row + 0.5, badge_d, badge_d, "F")

        # Number inside the badge
        pdf.set_xy(x_block, y_row + 0.5)
        pdf.set_font("Arial", "B", 8)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(badge_d, badge_d, str(n), align="C")

        # Name (left-aligned, same x for every row)
        pdf.set_xy(x_block + badge_d + gap, y_row)
        pdf.set_font("Arial", "", 13)
        pdf.set_text_color(*P["slate_900"])
        pdf.cell(max_name_w + 1, badge_d + 1, name, align="L")

        pdf.set_y(y_row + row_h)

    # Bottom: rose stripe + indigo band with repo URL
    rose_h = 4.0
    indigo_h = 14.0
    band_y = page_h - indigo_h - rose_h
    r, g, b = P["rose_600"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, band_y, page_w, rose_h, "F")
    r, g, b = P["indigo_900"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(0, band_y + rose_h, page_w, indigo_h, "F")

    if info.get("repository"):
        pdf.set_xy(0, band_y + rose_h + 4.5)
        pdf.set_font("Arial", "B", 9)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(page_w, 5, _clean_url(info["repository"]), align="C")

    pdf.set_text_color(0, 0, 0)


# ---------- PDF body ---------------------------------------------------------

def _render_pdf_h2_pill(pdf: PremiumPDF, text: str) -> None:
    """H2 as a filled rounded pill — indigo background, white text."""
    pdf.ln(4)
    pdf.set_font("Arial", "B", 13)
    text_w = pdf.get_string_width(text)
    pill_w = text_w + 12
    pill_h = 9.5
    # Atomic-fit: avoid stranding the pill at the very bottom of a page.
    if pdf.get_y() + pill_h + 4 > pdf.h - pdf.b_margin:
        pdf.add_page()
    x = pdf.l_margin
    y = pdf.get_y()
    r, g, b = P["indigo_700"]
    pdf.set_fill_color(r, g, b)
    pdf.set_draw_color(r, g, b)
    pdf.rect(x, y, pill_w, pill_h, "F")
    # Round-ish caps via small circles at each end
    pdf.ellipse(x - pill_h / 2 + 0.5, y, pill_h, pill_h, "F")
    pdf.ellipse(x + pill_w - pill_h / 2 - 0.5, y, pill_h, pill_h, "F")

    pdf.set_xy(x, y)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(pill_w, pill_h, text, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.set_y(y + pill_h + 3)


def _render_pdf_h3(pdf: PremiumPDF, text: str) -> None:
    pdf.ln(2)
    # Atomic-fit: estimate the wrapped text height before drawing the dot.
    avail_w = pdf.w - pdf.l_margin - pdf.r_margin - 5
    pdf.set_font("Arial", "B", 12)
    lines = pdf.multi_cell(avail_w, 6, text, dry_run=True, output="LINES")
    needed = max(7, len(lines) * 6) + 1
    if pdf.get_y() + needed > pdf.h - pdf.b_margin:
        pdf.add_page()
    y = pdf.get_y()
    # Small rose dot before the heading
    r, g, b = P["rose_600"]
    pdf.set_fill_color(r, g, b)
    pdf.ellipse(pdf.l_margin, y + 2.2, 2.4, 2.4, "F")

    pdf.set_xy(pdf.l_margin + 5, y)
    pdf.set_font("Arial", "B", 12)
    pdf.set_text_color(*P["indigo_900"])
    pdf.multi_cell(0, 6, text)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(0.5)


def _compute_col_widths(pdf, rows, n_cols, page_width):
    pdf.set_font("Arial", "", 10)
    raw = [0.0] * n_cols
    for row in rows:
        for c_idx in range(n_cols):
            cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
            for line in cell.splitlines() or [cell]:
                w = pdf.get_string_width(line)
                if w > raw[c_idx]:
                    raw[c_idx] = w

    pad = 5.0
    raw = [w + pad for w in raw]
    total = sum(raw) or 1.0
    widths = [max(15.0, w * page_width / total) for w in raw]
    scale = page_width / sum(widths)
    return [w * scale for w in widths]


def _draw_pdf_row(pdf, row, widths, col_align, h, line_h, v_pad,
                  is_header=False, zebra=False):
    n_cols = len(widths)
    y_start = pdf.get_y()
    for c_idx in range(n_cols):
        cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
        x_start = pdf.l_margin + sum(widths[:c_idx])
        if is_header:
            r, g, b = P["indigo_700"]
            pdf.set_fill_color(r, g, b)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Arial", "B", 10.5)
        else:
            if zebra:
                r, g, b = P["indigo_50"]
                pdf.set_fill_color(r, g, b)
            else:
                pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(*P["slate_900"])
            pdf.set_font("Arial", "", 10)
        r, g, b = P["indigo_100"]
        pdf.set_draw_color(r, g, b)
        pdf.set_line_width(0.2)

        align = "C"
        pdf.rect(x_start, y_start, widths[c_idx], h, "DF")

        inset_x = 1.6
        inset_y = v_pad / 2
        pdf.set_xy(x_start + inset_x, y_start + inset_y)
        pdf.multi_cell(widths[c_idx] - 2 * inset_x, line_h, cell,
                       border=0, align=align, fill=False)
    pdf.set_xy(pdf.l_margin, y_start + h)
    pdf.set_text_color(0, 0, 0)


def _render_table_pdf(pdf, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    widths = _compute_col_widths(pdf, rows, n_cols, page_width)

    col_align = ["C"] * n_cols  # every cell centered

    pdf.ln(2)
    line_h = 5.6
    v_pad = 2.4

    def measure(row, font_style):
        pdf.set_font("Arial", font_style, 10)
        max_lines = 1
        for c_idx in range(n_cols):
            cell = _strip_inline(row[c_idx]) if c_idx < len(row) else ""
            lines = pdf.multi_cell(widths[c_idx], line_h, cell,
                                   dry_run=True, output="LINES")
            if len(lines) > max_lines:
                max_lines = len(lines)
        return max_lines * line_h + v_pad

    for r_idx, row in enumerate(rows):
        is_header = r_idx == 0
        style = "B" if is_header else ""
        h = measure(row, style)
        if pdf.get_y() + h > pdf.h - pdf.b_margin:
            pdf.add_page()
            if not is_header and rows:
                hh = measure(rows[0], "B")
                _draw_pdf_row(pdf, rows[0], widths, col_align, hh, line_h, v_pad,
                              is_header=True)
        _draw_pdf_row(pdf, row, widths, col_align, h, line_h, v_pad,
                      is_header=is_header,
                      zebra=(r_idx % 2 == 0 and not is_header))
    pdf.ln(5)


def _render_pdf_quote(pdf, text: str) -> None:
    """Filled callout card with thick rose left border."""
    pdf.ln(2)
    y_start = pdf.get_y()
    # Pre-measure
    pdf.set_font("Arial", "I", 11)
    text_w = pdf.w - pdf.l_margin - pdf.r_margin - 8
    lines = pdf.multi_cell(text_w, 5.5, text, dry_run=True, output="LINES")
    box_h = max(8.0, len(lines) * 5.5 + 4.0)

    # Filled background
    r, g, b = P["indigo_50"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(pdf.l_margin, y_start, pdf.w - pdf.l_margin - pdf.r_margin, box_h, "F")
    # Left thick rose bar
    r, g, b = P["rose_600"]
    pdf.set_fill_color(r, g, b)
    pdf.rect(pdf.l_margin, y_start, 1.6, box_h, "F")

    pdf.set_xy(pdf.l_margin + 5, y_start + 2)
    pdf.set_font("Arial", "I", 11)
    pdf.set_text_color(*P["slate_700"])
    pdf.multi_cell(text_w, 5.5, text)
    pdf.set_text_color(0, 0, 0)
    pdf.set_y(y_start + box_h + 4)


def render_pdf(cover_blocks, body_blocks):
    pdf = PremiumPDF()
    # Cover page (no header/footer chrome)
    pdf._chrome = False
    pdf.add_page()

    if cover_blocks:
        info = _extract_cover_info(cover_blocks)
        _render_pdf_cover(pdf, info)

    # Body — enable running chrome
    pdf._chrome = True
    pdf.add_page()

    for kind, payload in body_blocks:
        if kind == "h1":
            pdf.set_font("Arial", "B", 22)
            pdf.set_text_color(*P["indigo_900"])
            pdf.multi_cell(0, 11, payload)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
        elif kind == "h2":
            _render_pdf_h2_pill(pdf, payload)
        elif kind == "h3":
            _render_pdf_h3(pdf, payload)
        elif kind == "p":
            write_justified_v2(pdf, payload)
            pdf.ln(2)
        elif kind == "ul":
            for item in payload:
                pdf.set_font("Arial", "", 11)
                pdf.cell(4, 5.4, "")
                pdf.set_text_color(*P["rose_600"])
                pdf.write(5.4, "■  ")
                pdf.set_text_color(0, 0, 0)
                pdf.add_inline(item)
                pdf.ln(5.6)
            pdf.ln(2)
        elif kind == "ol":
            for n, item in enumerate(payload, 1):
                pdf.set_font("Arial", "B", 11)
                pdf.cell(4, 5.4, "")
                pdf.set_text_color(*P["indigo_700"])
                pdf.write(5.4, f"{n:02d}.  ")
                pdf.set_text_color(0, 0, 0)
                pdf.add_inline(item)
                pdf.ln(5.6)
            pdf.ln(2)
        elif kind == "quote":
            _render_pdf_quote(pdf, payload)
        elif kind == "code":
            r, g, b = P["slate_100"]
            pdf.set_fill_color(r, g, b)
            pdf.set_font("Mono", "", 9)
            pdf.set_text_color(*P["slate_700"])
            for code_line in payload.split("\n"):
                pdf.cell(0, 4.6, code_line, ln=1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(3)
        elif kind == "math":
            pdf.set_font("Arial", "I", 11)
            pdf.set_text_color(*P["slate_700"])
            pdf.multi_cell(0, 5.5, payload)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(3)
        elif kind == "hr":
            y = pdf.get_y() + 2
            cx = pdf.w / 2
            r, g, b = P["rose_600"]
            pdf.set_draw_color(r, g, b)
            pdf.set_line_width(0.4)
            pdf.line(cx - 25, y, cx - 4, y)
            pdf.line(cx + 4, y, cx + 25, y)
            # Center diamond
            pdf.set_fill_color(r, g, b)
            pdf.ellipse(cx - 1.3, y - 1.3, 2.6, 2.6, "F")
            pdf.ln(6)
        elif kind == "table":
            _render_table_pdf(pdf, payload)

    pdf.output(str(PDF))
    print(f"wrote {PDF} ({PDF.stat().st_size/1024:.1f} KB)")


# ---------- Main -------------------------------------------------------------

def main():
    md_text = MD.read_text(encoding="utf-8")
    md_text = normalize_emoji(md_text)
    blocks = list(parse(md_text))
    cover_blocks, body_blocks = _split_cover(blocks)
    print(f"Parsed {len(blocks)} blocks (cover={len(cover_blocks)}, "
          f"body={len(body_blocks)}) from {MD}")

    build_docx(cover_blocks, body_blocks)
    render_pdf(cover_blocks, body_blocks)


if __name__ == "__main__":
    main()
